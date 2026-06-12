from __future__ import annotations

import json
import gzip
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import textwrap
import uuid
from datetime import datetime
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path
from urllib.parse import parse_qs, unquote, urlparse
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from backend.research_agent import benchmark_cases_status, build_client_report, build_model_oriented_research
except ModuleNotFoundError:
    from research_agent import benchmark_cases_status, build_client_report, build_model_oriented_research


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
FRONTEND_DIR = ROOT / "frontend"
EXPORT_DIR = ROOT / "exports"
PROJECTS_PATH = DATA_DIR / "projects.json"
LOCAL_IDENTITY_PATH = DATA_DIR / "local_identity.json"
ADMIN_STATION_OUTLINES_PATH = DATA_DIR / "admin_station_outlines.json"
STATION_MEMORY_PATH = DATA_DIR / "station_memory.json"
SCHEMATIC_DIR = FRONTEND_DIR / "schematic"
SCHEMATIC_GEOMETRY_PATH = SCHEMATIC_DIR / "user_geometry.json"
SCHEMATIC_EXPORT_DIR = SCHEMATIC_DIR / "exports"
GENERATED_IMAGE_DIR = EXPORT_DIR / "generated-images"


def load_json(name: str) -> dict:
    return json.loads((DATA_DIR / name).read_text(encoding="utf-8"))


def load_json_optional(name: str, default: dict) -> dict:
    path = DATA_DIR / name
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def amap_credentials() -> tuple[str, str]:
    key = os.environ.get("AMAP_JS_KEY", "")
    security_code = os.environ.get("AMAP_SECURITY_CODE", "")
    if key and security_code:
        return key, security_code
    try:
        import winreg

        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, "Environment") as env_key:
            if not key:
                key = winreg.QueryValueEx(env_key, "AMAP_JS_KEY")[0]
            if not security_code:
                security_code = winreg.QueryValueEx(env_key, "AMAP_SECURITY_CODE")[0]
    except Exception:
        pass
    return key, security_code


def subprocess_output_text(completed: subprocess.CompletedProcess, fallback: str) -> str:
    stderr = (completed.stderr or "").strip()
    stdout = (completed.stdout or "").strip()
    return stderr or stdout or fallback


FACTORS = load_json("factors.json")
RULES = load_json("design_rules.json")
STATIONS = load_json("stations.json")
DEMOS = load_json("demo_cases.json")
RIDERSHIP = load_json_optional("ridership.json", {"records": []})
RIDERSHIP_FORECAST = load_json_optional("ridership_forecast.json", {"records": []})
STATION_OPERATIONS = load_json_optional("station_operations.json", {"records": []})
STATION_AMENITIES = load_json_optional("station_amenities.json", {"records": []})
INPUT_SCHEMA = load_json_optional("input_schema.json", {"fields": []})
PPT_RULES = load_json_optional("ppt_rules_summary.json", {})
KNOWLEDGE_CATALOG = load_json_optional("knowledge/knowledge_catalog.json", {})
RULE_CARDS = load_json_optional("knowledge/rule_cards.json", {"cards": []})
STATION_KNOWLEDGE_INDEX = load_json_optional("knowledge/station_index.json", {"records": []})
SOURCE_MANIFEST = load_json_optional("knowledge/source_manifest.json", [])
UNPARSED_SOURCES = load_json_optional("knowledge/unparsed_sources.json", [])
STATION_INDEX = {item["name"]: item for item in STATIONS.get("stations", [])}


def source_manifest_records() -> list[dict]:
    records = list(SOURCE_MANIFEST if isinstance(SOURCE_MANIFEST, list) else [])
    if not (RIDERSHIP_FORECAST.get("records") or []):
        return records
    if any(item.get("id") == "ridership-forecast-0528-xls" for item in records):
        return records
    source = RIDERSHIP_FORECAST.get("source") or {}
    records.append({
        "id": "ridership-forecast-0528-xls",
        "path": "0528既有线路客流预测数据.xls",
        "name": source.get("fileName") or "0528既有线路客流预测数据.xls",
        "suffix": ".xls",
        "size": source.get("sizeBytes") or 0,
        "parseStatus": "parsed",
        "usefulness": "核心",
        "priority": 91,
        "category": "ridership",
        "reason": "既有线路全日客流量预测，用于站点上下文与报告中的未来客流证据，不替代现状日均进站数据",
        "counts": RIDERSHIP_FORECAST.get("counts") or {},
        "note": "解析说明见 docs/ridership_forecast_source.md；运行时使用 data/ridership_forecast.json",
    })
    return records


def utc_now() -> str:
    return datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def safe_slug(value: str | None, fallback: str = "project") -> str:
    text = value or fallback
    text = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", text, flags=re.UNICODE).strip("-")
    return text[:80] or fallback


def station_aliases(name: str | None) -> list[str]:
    text = (name or "").strip()
    if not text:
        return []
    aliases = [text]
    aliases.extend(match.strip() for match in re.findall(r"[（(]([^）)]+)[）)]", text) if match.strip())
    no_paren = re.sub(r"[（(].*?[）)]", "", text).strip()
    if no_paren:
        aliases.append(no_paren)
    if no_paren.endswith("站"):
        aliases.append(no_paren[:-1])
    return list(dict.fromkeys(item for item in aliases if item))


def build_alias_index(records: list[dict], name_getter) -> dict[str, dict]:
    index: dict[str, dict] = {}
    for record in records:
        names = name_getter(record)
        if isinstance(names, str):
            names = [names]
        for name in names or []:
            for alias in station_aliases(name):
                index.setdefault(alias, record)
    return index


def build_alias_multi_index(records: list[dict], name_getter) -> dict[str, list[dict]]:
    index: dict[str, list[dict]] = {}
    for record in records:
        names = name_getter(record)
        if isinstance(names, str):
            names = [names]
        for name in names or []:
            for alias in station_aliases(name):
                index.setdefault(alias, []).append(record)
    return index


RIDERSHIP_INDEX = build_alias_index(RIDERSHIP.get("records", []), lambda item: item.get("stationName", ""))
RIDERSHIP_FORECAST_INDEX = build_alias_multi_index(
    RIDERSHIP_FORECAST.get("records", []),
    lambda item: [item.get("stationName", ""), item.get("stationDisplayName", "")]
)
OPERATIONS_INDEX = build_alias_index(
    STATION_OPERATIONS.get("records", []),
    lambda item: [item.get("name", ""), *item.get("aliases", []), *item.get("displayNames", [])]
)
AMENITIES_INDEX = build_alias_index(
    STATION_AMENITIES.get("records", []),
    lambda item: [item.get("name", ""), *item.get("aliases", []), *item.get("displayNames", [])]
)
STATION_ALIAS_INDEX = build_alias_index(STATIONS.get("stations", []), lambda item: item.get("name", ""))
STATION_KNOWLEDGE_LOOKUP = build_alias_index(
    STATION_KNOWLEDGE_INDEX.get("records", []),
    lambda item: [item.get("name", ""), *item.get("aliases", [])]
)


def lookup_station_context(index: dict[str, dict], name: str | None) -> dict | None:
    for alias in station_aliases(name):
        if alias in index:
            return index[alias]
    return None


def lookup_station_context_list(index: dict[str, list[dict]], name: str | None) -> list[dict]:
    matches: list[dict] = []
    seen: set[int] = set()
    for alias in station_aliases(name):
        for record in index.get(alias, []):
            marker = id(record)
            if marker not in seen:
                seen.add(marker)
                matches.append(record)
    return matches


def load_projects() -> dict:
    if not PROJECTS_PATH.exists():
        return {"version": "1.0", "projects": []}
    return json.loads(PROJECTS_PATH.read_text(encoding="utf-8-sig"))


def save_projects(data: dict) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    PROJECTS_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def truthy_env(name: str) -> bool:
    return os.environ.get(name, "").strip().lower() in {"1", "true", "yes", "on"}


def env_first(*names: str, default: str = "") -> str:
    for name in names:
        value = os.environ.get(name)
        if value not in (None, ""):
            return str(value)
    return default


def int_env(*names: str, default: int) -> int:
    value = env_first(*names, default=str(default))
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def relative_to_root(path: Path) -> str:
    try:
        return path.relative_to(ROOT).as_posix()
    except ValueError:
        return str(path)


def server_config() -> dict:
    image_provider = env_first("GENERATED_IMAGE_PROVIDER", default="disabled").strip().lower() or "disabled"
    image_enabled = truthy_env("GENERATED_IMAGE_API_ENABLED") or image_provider not in {"disabled", "none", "off"}
    account_mode = env_first("INTERCONNECT_ACCOUNT_MODE", "ACCOUNT_MODE", default="local_anonymous").strip() or "local_anonymous"
    return {
        "host": env_first("INTERCONNECT_HOST", "HOST", default="0.0.0.0"),
        "port": int_env("INTERCONNECT_PORT", "PORT", default=8765),
        "dataDir": DATA_DIR,
        "exportDir": EXPORT_DIR,
        "stationMemoryPath": STATION_MEMORY_PATH,
        "localIdentityPath": LOCAL_IDENTITY_PATH,
        "secretKeyConfigured": bool(env_first("INTERCONNECT_SECRET_KEY", "SECRET_KEY")),
        "accountMode": account_mode,
        "accountsEnabled": account_mode not in {"", "local_anonymous", "anonymous", "disabled", "off"},
        "generatedImage": {
            "enabled": image_enabled,
            "provider": image_provider,
            "apiKeyConfigured": bool(os.environ.get("OPENAI_API_KEY")),
            "artifactDir": GENERATED_IMAGE_DIR,
        },
        "amap": {
            "jsKeyConfigured": bool(env_first("AMAP_JS_KEY")),
            "securityCodeConfigured": bool(env_first("AMAP_SECURITY_CODE")),
        },
        "exportPdfEnabled": export_pdf_enabled(),
    }


def path_writable(path: Path) -> bool:
    try:
        path.mkdir(parents=True, exist_ok=True)
        probe = path / ".write-test"
        probe.write_text("ok", encoding="utf-8")
        probe.unlink()
        return True
    except Exception:
        return False


def validate_server_configuration() -> dict:
    config = server_config()
    issues = []
    warnings = []
    for key in ("dataDir", "exportDir"):
        if not path_writable(config[key]):
            issues.append({"key": key, "message": f"{relative_to_root(config[key])} is not writable"})
    image = config["generatedImage"]
    if image["enabled"] and image["provider"] not in {"local", "local_placeholder", "mock"} and not image["apiKeyConfigured"]:
        issues.append({"key": "generatedImage", "message": "configured generated-image provider requires OPENAI_API_KEY or local provider"})
    if config["accountsEnabled"] and not config["secretKeyConfigured"]:
        warnings.append({"key": "accounts", "message": "account mode is enabled without INTERCONNECT_SECRET_KEY; use only for local validation"})
    return {
        "ok": not issues,
        "issues": issues,
        "warnings": warnings,
        "resolved": {
            "host": config["host"],
            "port": config["port"],
            "dataDir": relative_to_root(config["dataDir"]),
            "exportDir": relative_to_root(config["exportDir"]),
            "accountMode": config["accountMode"],
            "generatedImageProvider": config["generatedImage"]["provider"],
            "exportPdfEnabled": config["exportPdfEnabled"],
        },
    }


def build_platform_capability_status() -> dict:
    config = server_config()
    validation = validate_server_configuration()
    image = config["generatedImage"]
    image_configured = image["enabled"] and (
        image["provider"] in {"local", "local_placeholder", "mock"}
        or image["apiKeyConfigured"]
    )
    return {
        "generatedImage": {
            "enabled": image_configured,
            "mode": "configured" if image_configured else "not_configured",
            "endpoint": "/api/generated-images",
            "provider": image["provider"],
            "artifactDir": relative_to_root(image["artifactDir"]),
            "requires": ["GENERATED_IMAGE_PROVIDER=local"] if image["provider"] in {"disabled", "none", "off"} else ["OPENAI_API_KEY or local provider"],
        },
        "accounts": {
            "enabled": config["accountsEnabled"],
            "mode": config["accountMode"] if config["accountsEnabled"] else "local_anonymous",
            "identityEndpoint": "/api/identity",
            "ownerMode": config["accountMode"],
            "upgradePath": "Attach saved projects, exports, and schematic geometry to a future authenticated owner id.",
        },
        "adminStationOutlines": {
            "enabled": True,
            "mode": "shared_local_json",
            "storage": relative_to_root(STATION_MEMORY_PATH),
            "legacyStorage": relative_to_root(ADMIN_STATION_OUTLINES_PATH),
            "endpoints": ["/api/station-memory", "/api/station-memory/apply", "/api/admin/station-outlines", "/api/admin/station-outlines/apply"],
        },
        "deployment": {
            "enabled": True,
            "mode": "local_http_server",
            "host": config["host"],
            "port": config["port"],
            "dataDir": relative_to_root(config["dataDir"]),
            "exportDir": relative_to_root(config["exportDir"]),
            "exportPdfEnabled": config["exportPdfEnabled"],
            "secretKeyConfigured": config["secretKeyConfigured"],
            "validation": validation,
            "docs": "docs/deployment_server_migration.md",
        },
    }


def generated_image_placeholder(payload: dict | None = None) -> dict:
    return {
        "ok": False,
        "error": {
            "code": "not_configured",
            "message": "Generated-image API is not configured for this local delivery.",
            "required": ["GENERATED_IMAGE_API_ENABLED=1", "OPENAI_API_KEY"],
        },
        "request": payload or {},
        "capability": build_platform_capability_status()["generatedImage"],
    }


def xml_escape(value: str | None) -> str:
    return (
        str(value or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def generated_image_provider_failure(payload: dict, message: str) -> dict:
    return {
        "ok": False,
        "error": {
            "code": "provider_failure",
            "message": message,
        },
        "request": payload,
        "capability": build_platform_capability_status()["generatedImage"],
    }


def create_local_generated_image(payload: dict) -> dict:
    prompt = str(payload.get("prompt") or payload.get("description") or "station connection rendering").strip()
    GENERATED_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
    slug = safe_slug(prompt, "generated-image")
    image_path = GENERATED_IMAGE_DIR / f"{slug}-{stamp}.svg"
    metadata_path = GENERATED_IMAGE_DIR / f"{slug}-{stamp}.metadata.json"
    title = xml_escape(prompt[:80])
    subtitle = xml_escape(payload.get("stationName") or payload.get("projectName") or "Interconnect generated-image local provider")
    svg = f"""<svg xmlns="http://www.w3.org/2000/svg" width="1280" height="720" viewBox="0 0 1280 720">
  <rect width="1280" height="720" fill="#f7fbfd"/>
  <rect x="72" y="84" width="1136" height="552" rx="28" fill="#ffffff" stroke="#b8d4df" stroke-width="4"/>
  <path d="M170 470 C330 320 430 530 570 360 S820 270 1010 420" fill="none" stroke="#0f79b5" stroke-width="24" stroke-linecap="round"/>
  <path d="M170 500 C330 350 430 560 570 390 S820 300 1010 450" fill="none" stroke="#d1495b" stroke-width="10" stroke-linecap="round" stroke-dasharray="26 20"/>
  <circle cx="170" cy="470" r="42" fill="#0f79b5"/>
  <circle cx="1010" cy="420" r="42" fill="#d1495b"/>
  <text x="116" y="162" font-family="Microsoft YaHei, Arial, sans-serif" font-size="42" font-weight="700" fill="#16323f">{title}</text>
  <text x="116" y="220" font-family="Microsoft YaHei, Arial, sans-serif" font-size="24" fill="#5f7480">{subtitle}</text>
  <text x="116" y="594" font-family="Microsoft YaHei, Arial, sans-serif" font-size="22" fill="#5f7480">local generated-image provider artifact</text>
</svg>
"""
    image_path.write_text(svg, encoding="utf-8")
    metadata = {
        "provider": "local",
        "prompt": prompt,
        "createdAt": utc_now(),
        "owner": current_owner_metadata("generated_image"),
        "request": payload,
        "image": relative_to_root(image_path),
    }
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    return {
        "ok": True,
        "provider": "local",
        "image": file_info(image_path),
        "metadataFile": file_info(metadata_path),
        "metadata": metadata,
        "capability": build_platform_capability_status()["generatedImage"],
    }


def generated_image_response(payload: dict | None = None) -> dict:
    payload = payload or {}
    config = server_config()["generatedImage"]
    capability = build_platform_capability_status()["generatedImage"]
    if not capability["enabled"]:
        return generated_image_placeholder(payload)
    provider = config["provider"]
    if provider in {"local", "local_placeholder", "mock"}:
        return create_local_generated_image(payload)
    return generated_image_provider_failure(
        payload,
        f"Generated-image provider '{provider}' is configured but no runtime adapter is installed in this local delivery.",
    )


def local_identity_payload() -> dict:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    config = server_config()
    if LOCAL_IDENTITY_PATH.exists():
        data = json.loads(LOCAL_IDENTITY_PATH.read_text(encoding="utf-8-sig"))
    else:
        data = {
            "version": "local-anonymous.v1",
            "identity": {
                "id": f"anon-{uuid.uuid4().hex[:12]}",
                "type": "anonymous",
                "scope": "local-workstation",
                "createdAt": utc_now(),
            },
        }
        LOCAL_IDENTITY_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    if config["accountsEnabled"]:
        account_id = env_first("INTERCONNECT_ACCOUNT_ID", "ACCOUNT_ID", default="")
        existing_account = next((item for item in data.get("accounts", []) if item.get("type") == "local_user"), None)
        if not account_id:
            account_id = (existing_account or {}).get("id") or f"local-user-{uuid.uuid4().hex[:12]}"
        account = {
            "id": account_id,
            "type": "local_user",
            "scope": config["accountMode"],
            "displayName": env_first("INTERCONNECT_ACCOUNT_NAME", "ACCOUNT_NAME", default=(existing_account or {}).get("displayName") or "Local User"),
            "createdAt": (existing_account or data.get("identity", {})).get("createdAt") or utc_now(),
        }
        data["identity"] = account
        data.setdefault("accounts", [])
        if not any(item.get("id") == account_id for item in data["accounts"]):
            data["accounts"].append(account)
            LOCAL_IDENTITY_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    data["capability"] = build_platform_capability_status()["accounts"]
    return data


def current_owner_metadata(action: str = "local") -> dict:
    identity = local_identity_payload().get("identity") or {}
    config = server_config()
    return {
        "ownerId": identity.get("id") or "anonymous",
        "ownerType": identity.get("type") or "anonymous",
        "accountMode": config["accountMode"] if config["accountsEnabled"] else "local_anonymous",
        "scope": identity.get("scope") or "local-workstation",
        "action": action,
        "recordedAt": utc_now(),
    }


def migrate_owner_metadata(record: dict, target_owner: dict | None = None) -> dict:
    target = target_owner or current_owner_metadata("ownership_migration")
    previous = record.get("owner") or {"ownerId": "anonymous", "ownerType": "anonymous"}
    return {
        **target,
        "migration": {
            "from": previous,
            "migratedAt": utc_now(),
            "mode": "anonymous_to_account",
        },
    }


def load_admin_station_outline_data() -> dict:
    if not ADMIN_STATION_OUTLINES_PATH.exists():
        return {"version": "admin-station-outlines.v1", "records": []}
    return json.loads(ADMIN_STATION_OUTLINES_PATH.read_text(encoding="utf-8-sig"))


def save_admin_station_outline_data(data: dict) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    ADMIN_STATION_OUTLINES_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def list_admin_station_outlines(station_name: str | None = None) -> list[dict]:
    records = load_admin_station_outline_data().get("records", [])
    if station_name:
        records = [item for item in records if item.get("stationName") == station_name]
    return sorted(records, key=lambda item: item.get("updatedAt") or item.get("createdAt") or "", reverse=True)


def save_admin_station_outline(payload: dict) -> dict:
    station_name = (payload.get("stationName") or "").strip()
    outline = payload.get("outline") or {}
    path = outline.get("path") or outline.get("body")
    if not station_name:
        raise ValueError("stationName is required")
    if not isinstance(path, list) or len(path) < 3:
        raise ValueError("outline.path must include at least three points")
    data = load_admin_station_outline_data()
    record_id = payload.get("id") or outline.get("id") or uuid.uuid4().hex[:12]
    now = utc_now()
    existing = next((item for item in data.get("records", []) if item.get("id") == record_id), None)
    record = {
        "id": record_id,
        "stationName": station_name,
        "outline": {
            **outline,
            "id": outline.get("id") or f"admin-outline-{record_id}",
            "name": outline.get("name") or f"{station_name}站体轮廓",
            "path": path,
        },
        "source": payload.get("source") or {"type": "admin"},
        "createdAt": (existing or {}).get("createdAt") or now,
        "updatedAt": now,
    }
    data["records"] = [item for item in data.get("records", []) if item.get("id") != record_id]
    data["records"].append(record)
    save_admin_station_outline_data(data)
    return record


def apply_admin_station_outline_to_geometry(geometry: dict | None, station_name: str, outline_id: str | None = None) -> dict:
    geometry = json.loads(json.dumps(geometry or {}, ensure_ascii=False))
    candidates = list_admin_station_outlines(station_name)
    if outline_id:
        candidates = [item for item in candidates if item.get("id") == outline_id]
    if not candidates:
        return {"applied": False, "geometry": geometry, "error": "admin station outline not found"}
    record = candidates[0]
    outline = {
        **(record.get("outline") or {}),
        "source": {
            "kind": "admin_station_outline",
            "recordId": record.get("id"),
            "stationName": record.get("stationName"),
            "snapshotAt": utc_now(),
            "source": record.get("source") or {},
        },
    }
    existing = [
        item for item in geometry.get("stationOutlines", [])
        if ((item.get("source") or {}).get("recordId") != record.get("id"))
    ]
    geometry["stationOutlines"] = [outline, *existing]
    geometry["station"] = {
        **(geometry.get("station") or {}),
        "path": outline.get("path") or [],
        "body": outline.get("path") or [],
    }
    return {"applied": True, "record": record, "geometry": geometry}


def default_station_memory_data() -> dict:
    return {"version": "station-memory.v1", "records": []}


def load_station_memory_data() -> dict:
    if not STATION_MEMORY_PATH.exists():
        return default_station_memory_data()
    data = json.loads(STATION_MEMORY_PATH.read_text(encoding="utf-8-sig"))
    if not isinstance(data.get("records"), list):
        data["records"] = []
    data.setdefault("version", "station-memory.v1")
    return data


def save_station_memory_data(data: dict) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    data.setdefault("version", "station-memory.v1")
    data.setdefault("records", [])
    STATION_MEMORY_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def station_memory_identity(record: dict) -> dict:
    return record.get("identity") or {}


def station_memory_context(record: dict) -> dict:
    return record.get("context") or {}


def station_memory_alias_values(record: dict) -> list[str]:
    identity = station_memory_identity(record)
    return [
        identity.get("canonicalName") or "",
        identity.get("displayName") or "",
        *(identity.get("aliases") or []),
    ]


def station_memory_alias_set(record: dict) -> set[str]:
    aliases: set[str] = set()
    for value in station_memory_alias_values(record):
        aliases.update(station_aliases(value))
    return aliases


def normalize_station_memory_record(payload: dict, existing: dict | None = None) -> dict:
    project = payload.get("project") or {}
    station = project.get("station") or payload.get("station") or {}
    identity_payload = payload.get("identity") or {}
    context_payload = payload.get("context") or {}
    schematic_payload = payload.get("schematic") or {}
    source_labels = payload.get("sourceLabels") or []
    now = utc_now()

    canonical_name = first_present(
        identity_payload.get("canonicalName"),
        payload.get("stationName"),
        station.get("name"),
        (existing or {}).get("identity", {}).get("canonicalName"),
    )
    if not canonical_name:
        raise ValueError("stationName or identity.canonicalName is required")
    display_name = first_present(identity_payload.get("displayName"), payload.get("displayName"), canonical_name)
    aliases = [
        canonical_name,
        display_name,
        *(identity_payload.get("aliases") or []),
        *(payload.get("aliases") or []),
    ]
    aliases = list(dict.fromkeys(str(item).strip() for item in aliases if str(item or "").strip()))

    context = {
        "line": first_present(context_payload.get("line"), station.get("line"), (existing or {}).get("context", {}).get("line")) or "",
        "todLevel": first_present(context_payload.get("todLevel"), station.get("todLevel"), (existing or {}).get("context", {}).get("todLevel")) or "",
        "locationLevel": first_present(context_payload.get("locationLevel"), station.get("locationLevel"), (existing or {}).get("context", {}).get("locationLevel")) or "",
        "stationType": first_present(context_payload.get("stationType"), station.get("stationType"), (existing or {}).get("context", {}).get("stationType")) or "",
        "district": first_present(context_payload.get("district"), station.get("district"), (existing or {}).get("context", {}).get("district")) or "",
        "dailyInbound": first_present(context_payload.get("dailyInbound"), station.get("dailyInbound"), (existing or {}).get("context", {}).get("dailyInbound")),
        "nearbyExit": first_present(context_payload.get("nearbyExit"), station.get("nearbyExit"), (existing or {}).get("context", {}).get("nearbyExit")) or "",
        "interfaceCondition": first_present(context_payload.get("interfaceCondition"), station.get("interfaceCondition"), (existing or {}).get("context", {}).get("interfaceCondition")) or "",
    }

    schematic = {
        "stationOutlines": schematic_payload.get("stationOutlines") or (existing or {}).get("schematic", {}).get("stationOutlines") or [],
        "exits": schematic_payload.get("exits") or (existing or {}).get("schematic", {}).get("exits") or [],
        "assets": schematic_payload.get("assets") or (existing or {}).get("schematic", {}).get("assets") or [],
    }

    record_id = payload.get("id") or (existing or {}).get("id") or f"station-memory-{safe_slug(canonical_name, 'station')}-{uuid.uuid4().hex[:8]}"
    version = int((existing or {}).get("version") or 0) + 1
    return {
        "id": record_id,
        "version": version,
        "identity": {
            "canonicalName": canonical_name,
            "displayName": display_name,
            "aliases": aliases,
        },
        "context": context,
        "schematic": schematic,
        "notes": payload.get("notes") or (existing or {}).get("notes") or [],
        "fieldSources": payload.get("fieldSources") or (existing or {}).get("fieldSources") or {},
        "owner": payload.get("owner") or (existing or {}).get("owner") or current_owner_metadata("station_memory"),
        "provenance": {
            **((existing or {}).get("provenance") or {}),
            **(payload.get("provenance") or {}),
            "sourceLabels": source_labels,
            "createdFromProjectId": payload.get("createdFromProjectId") or project.get("id") or (existing or {}).get("provenance", {}).get("createdFromProjectId") or "",
            "operatorIntent": payload.get("operatorIntent") or "save_station_memory",
        },
        "createdAt": (existing or {}).get("createdAt") or now,
        "updatedAt": now,
    }


def admin_outline_memory_records() -> list[dict]:
    records = []
    for item in load_admin_station_outline_data().get("records", []):
        station_name = item.get("stationName") or ""
        outline = item.get("outline") or {}
        records.append({
            "id": f"memory-admin-{item.get('id')}",
            "version": 1,
            "virtual": True,
            "identity": {
                "canonicalName": station_name,
                "displayName": station_name,
                "aliases": station_aliases(station_name),
            },
            "context": {},
            "schematic": {
                "stationOutlines": [{
                    **outline,
                    "source": {
                        "kind": "admin_station_outline",
                        "recordId": item.get("id"),
                        "stationName": station_name,
                        "source": item.get("source") or {},
                    },
                }],
                "exits": [],
                "assets": [],
            },
            "notes": [],
            "fieldSources": {},
            "owner": {
                "ownerId": "administrator",
                "ownerType": "administrator",
                "accountMode": "shared_local_json",
                "scope": "station-memory",
                "action": "admin_station_outline",
            },
            "provenance": {
                "sourceLabels": ["administrator station outline"],
                "createdFromProjectId": "",
                "operatorIntent": "mapped_admin_station_outline",
            },
            "createdAt": item.get("createdAt") or "",
            "updatedAt": item.get("updatedAt") or "",
        })
    return records


def list_station_memory_records(station_name: str | None = None) -> list[dict]:
    records = list(load_station_memory_data().get("records", []))
    existing_ids = {item.get("id") for item in records}
    records.extend(item for item in admin_outline_memory_records() if item.get("id") not in existing_ids)
    if station_name:
        query_aliases = set(station_aliases(station_name))
        records = [
            item for item in records
            if query_aliases.intersection(station_memory_alias_set(item))
        ]
    return sorted(records, key=lambda item: item.get("updatedAt") or item.get("createdAt") or "", reverse=True)


def lookup_station_memory_record(name: str | None) -> dict | None:
    aliases = set(station_aliases(name))
    for record in list_station_memory_records():
        if aliases.intersection(station_memory_alias_set(record)):
            return record
    return None


def save_station_memory_record(payload: dict) -> dict:
    data = load_station_memory_data()
    record_id = payload.get("id")
    existing = next((item for item in data.get("records", []) if item.get("id") == record_id), None)
    if not existing:
        lookup_name = payload.get("stationName") or (payload.get("identity") or {}).get("canonicalName")
        existing = next((
            item for item in data.get("records", [])
            if lookup_name and set(station_aliases(lookup_name)).intersection(station_memory_alias_set(item))
        ), None)
    record = normalize_station_memory_record(payload, existing)
    data["records"] = [item for item in data.get("records", []) if item.get("id") != record["id"]]
    data["records"].append(record)
    save_station_memory_data(data)
    return record


def station_memory_snapshot(record: dict, applied_fields: list[str]) -> dict:
    context = station_memory_context(record)
    schematic = record.get("schematic") or {}
    return {
        "sourceMemoryId": record.get("id"),
        "sourceVersion": record.get("version") or 1,
        "appliedAt": utc_now(),
        "appliedFields": applied_fields,
        "context": context,
        "schematicSummary": {
            "stationOutlineCount": len(schematic.get("stationOutlines") or []),
            "exitCount": len(schematic.get("exits") or []),
            "assetCount": len(schematic.get("assets") or []),
        },
    }


def apply_station_memory_to_project(project: dict | None, record: dict, force: bool = True) -> dict:
    project = json.loads(json.dumps(project or {}, ensure_ascii=False))
    station = project.setdefault("station", {})
    identity = station_memory_identity(record)
    context = station_memory_context(record)
    applied_fields = []
    field_map = {
        "name": identity.get("canonicalName"),
        "line": context.get("line"),
        "todLevel": context.get("todLevel"),
        "locationLevel": context.get("locationLevel"),
        "stationType": context.get("stationType"),
        "district": context.get("district"),
        "dailyInbound": context.get("dailyInbound"),
        "nearbyExit": context.get("nearbyExit"),
        "interfaceCondition": context.get("interfaceCondition"),
    }
    for key, value in field_map.items():
        if value in (None, "", []):
            continue
        if force or station.get(key) in (None, "", []):
            station[key] = value
            applied_fields.append(f"station.{key}")
    project["stationMemorySnapshot"] = station_memory_snapshot(record, applied_fields)
    return project


def apply_station_memory_to_geometry(geometry: dict | None, record: dict) -> dict:
    geometry = json.loads(json.dumps(geometry or {}, ensure_ascii=False))
    schematic = record.get("schematic") or {}
    outlines = []
    for outline in schematic.get("stationOutlines") or []:
        outlines.append({
            **outline,
            "source": {
                **(outline.get("source") or {}),
                "kind": (outline.get("source") or {}).get("kind") or "station_memory",
                "memoryId": record.get("id"),
                "memoryVersion": record.get("version") or 1,
                "snapshotAt": utc_now(),
            },
        })
    if outlines:
        existing = geometry.get("stationOutlines") or []
        geometry["stationOutlines"] = [*outlines, *[
            item for item in existing
            if (item.get("source") or {}).get("memoryId") != record.get("id")
        ]]
        first = outlines[0]
        geometry["station"] = {
            **(geometry.get("station") or {}),
            "path": first.get("path") or [],
            "body": first.get("path") or [],
        }
    return geometry


def apply_station_memory_payload(payload: dict) -> dict:
    memory_id = payload.get("memoryId") or payload.get("id")
    station_name = payload.get("stationName") or ((payload.get("project") or {}).get("station") or {}).get("name")
    records = list_station_memory_records(station_name)
    if memory_id:
        records = [item for item in list_station_memory_records() if item.get("id") == memory_id]
    if not records:
        raise ValueError("station memory record not found")
    record = records[0]
    force = bool(payload.get("force", True))
    project = apply_station_memory_to_project(payload.get("project") or {}, record, force)
    geometry = apply_station_memory_to_geometry(payload.get("geometry") or {}, record)
    return {
        "ok": True,
        "record": record,
        "project": project,
        "geometry": geometry,
        "snapshot": project.get("stationMemorySnapshot"),
    }


def list_project_summaries() -> list[dict]:
    data = load_projects()
    summaries = []
    for item in data.get("projects", []):
        project = item.get("project", {})
        result = item.get("result") or {}
        if not isinstance(result, dict):
            result = {}
        summaries.append({
            "id": item.get("id"),
            "name": project.get("name") or "未命名项目",
            "projectCode": project.get("projectCode") or "",
            "stationName": (project.get("station") or {}).get("name") or "",
            "level": result.get("level") or "",
            "score": result.get("score"),
            "recommendedType": ((result.get("recommendation") or {}).get("primary") or {}).get("name") or "",
            "owner": item.get("owner") or project.get("owner") or {},
            "updatedAt": item.get("updatedAt"),
            "createdAt": item.get("createdAt")
        })
    return sorted(summaries, key=lambda item: item.get("updatedAt") or "", reverse=True)


def score_overview_rows(result: dict) -> list[dict]:
    rows = []
    for dimension in result.get("dimensions") or []:
        factors = dimension.get("factors") or []
        key_factors = []
        source_basis = []
        assumed_names = []
        for factor in factors:
            key_factors.append(
                f"{factor.get('name') or ''}：{factor.get('label') or '待补齐'}"
                f"（{float(factor.get('weightedScore') or 0):.4f}）"
            )
            source_basis.append(
                f"{factor.get('name') or ''}：{factor.get('source') or '规则评分表/项目输入'}"
            )
            if factor.get("assumed"):
                assumed_names.append(factor.get("name") or "")
        rows.append({
            "dimension": dimension.get("name") or "",
            "score": float(dimension.get("score") or 0),
            "keyFactors": "；".join(key_factors),
            "sourceBasis": "；".join(source_basis),
            "note": f"含临时试算：{'、'.join(assumed_names)}" if assumed_names else "核心字段已取值"
        })
    return rows


def option_lookup(factor: dict) -> dict:
    return {option["value"]: option for option in factor.get("options", [])}


def iter_factors() -> list[tuple[dict, dict]]:
    pairs: list[tuple[dict, dict]] = []
    for dimension in FACTORS["dimensions"]:
        for factor in dimension["factors"]:
            pairs.append((dimension, factor))
    return pairs


FACTOR_LOOKUP = {factor["id"]: option_lookup(factor) for _, factor in iter_factors()}


def classify_development_intensity(building_area: float | int | None) -> str | None:
    if building_area is None:
        return None
    try:
        value = float(building_area)
    except (TypeError, ValueError):
        return None
    if value >= 100000:
        return "gt_100k"
    if value >= 30000:
        return "30k_100k"
    if value >= 5000:
        return "5k_30k"
    return "lt_5k"


def classify_ridership(daily_inbound: float | int | None) -> str | None:
    if daily_inbound is None:
        return None
    try:
        value = float(daily_inbound)
    except (TypeError, ValueError):
        return None
    if value >= 9000:
        return "gt_9000"
    if value >= 5000:
        return "5000_9000"
    if value >= 2500:
        return "2500_5000"
    return "lt_2500"


def station_forecast_records(station_name: str | None, line: str | None = None) -> list[dict]:
    records = lookup_station_context_list(RIDERSHIP_FORECAST_INDEX, station_name)
    requested_lines = {
        item.lstrip("0")
        for item in re.split(r"[/、,，\s]+", str(line or ""))
        if item.strip()
    }
    if requested_lines:
        filtered = [
            record for record in records
            if str(record.get("line") or "").lstrip("0") in requested_lines
        ]
        if filtered:
            return filtered
    return records


def forecast_station_rollup(records: list[dict]) -> dict:
    horizons: dict[str, dict] = {}
    source = RIDERSHIP_FORECAST.get("source") or {}
    source_name = source.get("fileName") if isinstance(source, dict) else str(source or "")
    source_name = source_name or "0528既有线路客流预测数据.xls"
    lines = sorted({
        str(record.get("line") or "").lstrip("0") or str(record.get("line") or "")
        for record in records
        if record.get("line") not in (None, "")
    })
    for record in records:
        year = str(record.get("horizonYear") or record.get("year") or "")
        if not year:
            continue
        bucket = horizons.setdefault(year, {
            "horizonYear": int(year) if year.isdigit() else year,
            "boardingTotal": 0,
            "alightingTotal": 0,
            "directions": [],
        })
        boarding = numeric_value(record.get("boarding"))
        alighting = numeric_value(record.get("alighting"))
        bucket["boardingTotal"] += boarding
        bucket["alightingTotal"] += alighting
        bucket["directions"].append({
            "line": record.get("line"),
            "directionLabel": record.get("directionLabel") or record.get("direction") or "",
            "boarding": boarding,
            "alighting": alighting,
        })
    rollups = []
    for item in horizons.values():
        item["boardingTotal"] = int(round(item["boardingTotal"]))
        item["alightingTotal"] = int(round(item["alightingTotal"]))
        rollups.append(item)
    rollups.sort(key=lambda item: str(item.get("horizonYear")))
    return {
        "source": source_name,
        "sourceMetadata": source if isinstance(source, dict) else {},
        "unit": RIDERSHIP_FORECAST.get("unit") or "人次",
        "count": len(records),
        "lines": lines,
        "horizons": rollups,
        "records": records,
    }


def resolve_station_context(station: dict) -> dict:
    station_name = station.get("name")
    memory_record = lookup_station_memory_record(station_name)
    memory_context = (memory_record or {}).get("context") or {}
    ridership_record = lookup_station_context(RIDERSHIP_INDEX, station_name)
    operations_record = lookup_station_context(OPERATIONS_INDEX, station_name)
    amenities_record = lookup_station_context(AMENITIES_INDEX, station_name)
    forecast_records = station_forecast_records(station_name, station.get("line"))

    input_daily = station.get("dailyInbound")
    try:
        input_daily_value = float(input_daily) if input_daily not in {None, ""} else None
    except (TypeError, ValueError):
        input_daily_value = None

    matched_latest = ridership_record.get("latestDailyInbound") if ridership_record else None
    if input_daily_value is not None:
        daily_inbound = input_daily_value
        if matched_latest is not None and abs(input_daily_value - float(matched_latest)) < 1:
            daily_source = f"每站每月日均进站.xlsx：{ridership_record.get('latestMonth') or '最新月'}日均进站"
        else:
            daily_source = "使用方输入：日均客流"
    elif memory_context.get("dailyInbound") not in (None, ""):
        try:
            daily_inbound = float(memory_context["dailyInbound"])
            daily_source = "station memory"
        except (TypeError, ValueError):
            daily_inbound = None
            daily_source = ""
    elif ridership_record and ridership_record.get("latestDailyInbound") is not None:
        daily_inbound = float(ridership_record["latestDailyInbound"])
        daily_source = f"每站每月日均进站.xlsx：{ridership_record.get('latestMonth') or '最新月'}日均进站"
    else:
        daily_inbound = None
        daily_source = ""

    return {
        "dailyInbound": daily_inbound,
        "dailyInboundSource": daily_source,
        "memory": memory_record,
        "ridership": ridership_record,
        "ridershipForecast": forecast_station_rollup(forecast_records) if forecast_records else None,
        "operations": operations_record,
        "amenities": amenities_record
    }


def max_weighted_score() -> float:
    configured = (FACTORS.get("scoreScale") or {}).get("rawMax")
    if configured:
        return float(configured)
    total = 0.0
    for _, factor in iter_factors():
        total += factor["weight"] * max(option["score"] for option in factor.get("options", []))
    return round(total, 4)


def normalized_score(score: float) -> float:
    max_score = max_weighted_score()
    if max_score <= 0:
        return 0.0
    return round(score / max_score * 100, 2)


def compact_nearby_places(amenities: dict) -> str:
    nearby = amenities.get("nearby") or {}
    labels = {
        "schools": "学校",
        "hubs": "交通/公益",
        "residential": "住宅",
        "commercial": "商业景点",
        "other": "其他"
    }
    parts = []
    for key, label in labels.items():
        values = [item for item in nearby.get(key, []) if item]
        if values:
            parts.append(f"{label}：{'、'.join(values[:2])}")
    return "；".join(parts)


def load_knowledge_chunks() -> list[dict]:
    path = DATA_DIR / "knowledge" / "knowledge_chunks.jsonl"
    if not path.exists():
        return []
    chunks = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.strip():
            chunks.append(json.loads(line))
    return chunks


def search_knowledge(
    query: str,
    limit: int = 20,
    category: str | None = None,
    kind: str | None = None,
    min_score: int | None = None
) -> dict:
    terms = [term.lower() for term in re.split(r"\s+", query.strip()) if term.strip()]
    if not terms:
        return {"query": query, "count": 0, "results": []}
    results = []
    for chunk in load_knowledge_chunks():
        haystack = f"{chunk.get('title', '')} {chunk.get('text', '')}".lower()
        if not all(term in haystack for term in terms):
            continue
        hits = sum(haystack.count(term) for term in terms)
        if hits:
            score = hits * 10 + len(terms) * 100 + int(chunk.get("priority") or 0)
            if category and chunk.get("category") != category:
                continue
            if kind and chunk.get("kind") != kind:
                continue
            if min_score is not None and score < min_score:
                continue
            results.append({
                "score": score,
                "sourcePath": chunk.get("sourcePath"),
                "sourceName": chunk.get("sourceName"),
                "title": chunk.get("title"),
                "category": chunk.get("category"),
                "kind": chunk.get("kind"),
                "text": chunk.get("text", "")[:700],
                "metadata": chunk.get("metadata") or {}
            })
    results.sort(key=lambda item: item["score"], reverse=True)
    return {"query": query, "count": len(results), "results": results[:limit]}


def infer_transfer(station: dict, preset: dict | None) -> str | None:
    station_type = station.get("stationType")
    if station_type in {"current_transfer", "planned_transfer", "normal"}:
        return station_type
    lines = str(station.get("line") or (preset or {}).get("lines") or "")
    if parse_line_count(lines) >= 2:
        return "current_transfer"
    if station.get("name") or preset:
        return "normal"
    return None


def infer_location(station: dict, preset: dict | None) -> str | None:
    explicit = station.get("locationLevel")
    if explicit in FACTOR_LOOKUP["location_level"]:
        return explicit
    preset_level = (preset or {}).get("locationLevel")
    if preset_level in FACTOR_LOOKUP["location_level"]:
        return preset_level
    return None


def grade_score(score: float, score_percent: float | None = None) -> dict:
    for grade in FACTORS["grading"]:
        if score_percent is not None and ("percentMin" in grade or "percentMax" in grade):
            min_value = grade.get("percentMin")
            max_value = grade.get("percentMax")
            value = score_percent
        else:
            min_value = grade["min"]
            max_value = grade["max"]
            value = score
        if (min_value is None or value >= min_value) and (max_value is None or value < max_value):
            return grade
    return FACTORS["grading"][-1]


def parse_line_count(value: str | None) -> int:
    parts = [item for item in re.split(r"[/、,，\s]+", str(value or "")) if item]
    return len(parts)


def first_present(*values):
    for value in values:
        if value not in (None, "", []):
            return value
    return None


def join_values(value) -> str:
    if isinstance(value, list):
        return "/".join(str(item).strip() for item in value if str(item).strip())
    return str(value or "").strip()


def station_context_records(name: str | None) -> dict:
    forecast_records = station_forecast_records(name)
    return {
        "memory": lookup_station_memory_record(name),
        "station": lookup_station_context(STATION_ALIAS_INDEX, name),
        "ridership": lookup_station_context(RIDERSHIP_INDEX, name),
        "ridershipForecast": forecast_station_rollup(forecast_records) if forecast_records else None,
        "operations": lookup_station_context(OPERATIONS_INDEX, name),
        "amenities": lookup_station_context(AMENITIES_INDEX, name),
        "knowledge": lookup_station_context(STATION_KNOWLEDGE_LOOKUP, name),
    }


def canonical_station_name(name: str | None, records: dict) -> str:
    memory = records.get("memory") or {}
    memory_identity = memory.get("identity") or {}
    station = records.get("station") or {}
    ridership = records.get("ridership") or {}
    forecast = records.get("ridershipForecast") or {}
    operations = records.get("operations") or {}
    amenities = records.get("amenities") or {}
    knowledge = records.get("knowledge") or {}
    forecast_record = ((forecast.get("records") or [{}])[0] or {})
    return first_present(
        station.get("name"),
        memory_identity.get("canonicalName"),
        memory_identity.get("displayName"),
        ridership.get("stationName"),
        forecast_record.get("stationName"),
        knowledge.get("name"),
        operations.get("name"),
        amenities.get("name"),
        name,
    ) or ""


def station_line_text(records: dict, overrides: dict | None = None) -> str:
    overrides = overrides or {}
    memory = records.get("memory") or {}
    memory_context = memory.get("context") or {}
    station = records.get("station") or {}
    ridership = records.get("ridership") or {}
    forecast = records.get("ridershipForecast") or {}
    operations = records.get("operations") or {}
    amenities = records.get("amenities") or {}
    forecast_lines = forecast.get("lines") or []
    return join_values(first_present(
        overrides.get("line"),
        memory_context.get("line"),
        station.get("lines"),
        ridership.get("lines"),
        forecast_lines,
        operations.get("lines"),
        amenities.get("lines"),
    ))


def station_interface_summary(operations: dict | None, amenities: dict | None) -> str:
    operations = operations or {}
    amenities = amenities or {}
    parts = []
    if operations:
        forms = "、".join(operations.get("connectionForms") or [])
        parts.append(
            f"已识别出入口{operations.get('exitCount') or 0}个"
            f"；接口{operations.get('interfaceCount') or 0}个"
            f"；联通形式：{forms or '暂无登记'}"
        )
    if amenities:
        parts.append(
            f"开放出入口{amenities.get('openExitCount') or 0}/{amenities.get('exitRows') or 0}个"
            f"；运营管理{amenities.get('managedExitCount') or 0}个"
        )
    return "；".join(parts)


def station_sources(records: dict) -> list[dict]:
    labels = {
        "memory": "station memory",
        "station": "TOD station preset",
        "ridership": "ridership workbook",
        "ridershipForecast": "ridership forecast workbook",
        "operations": "station interface workbook",
        "amenities": "station amenity workbook",
        "knowledge": "station knowledge index",
    }
    return [
        {"key": key, "label": label, "matched": bool(records.get(key))}
        for key, label in labels.items()
        if records.get(key)
    ]


def station_context_payload(name: str | None, overrides: dict | None = None) -> dict:
    records = station_context_records(name)
    resolved_name = canonical_station_name(name, records)
    memory = records.get("memory") or {}
    memory_context = memory.get("context") or {}
    station = records.get("station") or {}
    ridership = records.get("ridership") or {}
    forecast = records.get("ridershipForecast") or {}
    operations = records.get("operations") or {}
    amenities = records.get("amenities") or {}
    overrides = overrides or {}
    line = station_line_text(records, overrides)
    forecast_records = station_forecast_records(resolved_name or name, line)
    if forecast_records:
        records["ridershipForecast"] = forecast_station_rollup(forecast_records)
        forecast = records["ridershipForecast"]
    context_input = {"name": resolved_name, **overrides}
    if line and not context_input.get("line"):
        context_input["line"] = line
    if memory_context.get("stationType") and not context_input.get("stationType"):
        context_input["stationType"] = memory_context.get("stationType")
    station_type = infer_transfer(context_input, station)
    suggested_fields = {
        "station.name": resolved_name,
        "station.line": line,
        "station.todLevel": first_present(overrides.get("todLevel"), memory_context.get("todLevel"), station.get("todLevel")) or "",
        "station.locationLevel": first_present(overrides.get("locationLevel"), memory_context.get("locationLevel"), station.get("locationLevel")) or "",
        "station.stationType": station_type or "",
        "station.dailyInbound": first_present(overrides.get("dailyInbound"), memory_context.get("dailyInbound"), ridership.get("latestDailyInbound")) or "",
        "station.district": first_present(
            overrides.get("district"),
            memory_context.get("district"),
            (operations.get("districts") or [None])[0],
            amenities.get("district"),
        ) or "",
        "station.nearbyExit": first_present(
            overrides.get("nearbyExit"),
            memory_context.get("nearbyExit"),
            ((amenities.get("sampleExits") or [{}])[0] or {}).get("exit"),
        ) or "",
        "station.interfaceCondition": first_present(
            overrides.get("interfaceCondition"),
            memory_context.get("interfaceCondition"),
            station_interface_summary(operations, amenities),
        ) or "",
    }
    return {
        "ok": any(records.values()),
        "query": name or "",
        "name": resolved_name,
        "stationType": station_type,
        "suggestedFields": suggested_fields,
        "sources": station_sources(records),
        **records,
    }


def iter_station_candidate_names() -> list[str]:
    names: list[str] = []
    for record in list_station_memory_records():
        names.extend(station_memory_alias_values(record))
    for station in STATIONS.get("stations", []):
        names.append(station.get("name", ""))
    for record in RIDERSHIP.get("records", []):
        names.append(record.get("stationName", ""))
    for record in RIDERSHIP_FORECAST.get("records", []):
        names.extend([record.get("stationName", ""), record.get("stationDisplayName", "")])
    for record in STATION_OPERATIONS.get("records", []):
        names.extend([record.get("name", ""), *record.get("aliases", []), *record.get("displayNames", [])])
    for record in STATION_AMENITIES.get("records", []):
        names.extend([record.get("name", ""), *record.get("aliases", []), *record.get("displayNames", [])])
    for record in STATION_KNOWLEDGE_INDEX.get("records", []):
        names.extend([record.get("name", ""), *record.get("aliases", [])])
    seen = set()
    result = []
    for name in names:
        canonical = canonical_station_name(name, station_context_records(name))
        if canonical and canonical not in seen:
            seen.add(canonical)
            result.append(canonical)
    return result


def station_match_score(query: str, name: str, context: dict) -> int:
    aliases = station_aliases(name)
    memory = context.get("memory") or {}
    aliases.extend(station_memory_alias_values(memory))
    knowledge = context.get("knowledge") or {}
    aliases.extend(knowledge.get("aliases") or [])
    normalized_query = query.strip().lower()
    if not normalized_query:
        return 10 + len(context.get("sources") or [])
    score = 0
    for alias in aliases:
        candidate = str(alias or "").strip().lower()
        if candidate == normalized_query:
            score = max(score, 1000)
        elif candidate.startswith(normalized_query):
            score = max(score, 760)
        elif normalized_query in candidate:
            score = max(score, 520)
    if score:
        score += len(context.get("sources") or []) * 20
        if context.get("memory"):
            score += 80
        if context.get("station"):
            score += 30
    return score


def search_stations(query: str, limit: int = 12) -> dict:
    query = (query or "").strip()
    results = []
    for name in iter_station_candidate_names():
        context = station_context_payload(name)
        score = station_match_score(query, name, context)
        if not score:
            continue
        fields = context.get("suggestedFields") or {}
        source_labels = [item["label"] for item in context.get("sources") or []]
        results.append({
            "name": context.get("name") or name,
            "label": " / ".join(item for item in [
                fields.get("station.line"),
                fields.get("station.todLevel"),
                fields.get("station.locationLevel"),
            ] if item),
            "line": fields.get("station.line") or "",
            "todLevel": fields.get("station.todLevel") or "",
            "locationLevel": fields.get("station.locationLevel") or "",
            "stationType": fields.get("station.stationType") or "",
            "sourceLabels": source_labels,
            "matchScore": score,
        })
    results.sort(key=lambda item: (item["matchScore"], len(item["sourceLabels"]), item["name"]), reverse=True)
    limit = min(max(int(limit or 12), 1), 50)
    return {"query": query, "count": len(results), "results": results[:limit]}


def numeric_value(value) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def pick_station_precheck_rule(station: dict | None, station_context: dict | None) -> str:
    station = station or {}
    station_context = station_context or {}
    operations = station_context.get("operations") or {}
    ridership = station_context.get("ridership") or {}
    forms = " ".join(operations.get("connectionForms") or [])
    inbound = numeric_value(station.get("dailyInbound") or ridership.get("latestDailyInbound"))
    line_count = parse_line_count(station.get("line") or ridership.get("lines"))
    location_level = station.get("locationLevel") or ""
    exit_count = int(operations.get("exitCount") or 0)

    if any(keyword in forms for keyword in ("高架", "天桥", "空中")):
        return "station_precheck_sky_bridge"
    if line_count >= 2 and inbound >= 9000 and location_level == "city":
        return "station_precheck_underground_main"
    if (line_count >= 2 and inbound >= 5000) or location_level == "city" or inbound >= 12000:
        return "station_precheck_underground_secondary"
    if location_level == "district" or inbound >= 5000 or exit_count >= 4:
        return "station_precheck_weather_corridor"
    return "fallback_surface"


def pick_rule(level: str, parcel: dict, missing: list[dict], station: dict | None = None, station_context: dict | None = None) -> dict:
    land_use = parcel.get("landUse")
    underground = parcel.get("undergroundSpace")
    has_critical_missing = any(item["factorId"] in {"land_use", "underground_space"} for item in missing)

    if has_critical_missing:
        rule_id = pick_station_precheck_rule(station, station_context)
    elif level == "必连地块" and underground == "commercial_public_pr":
        rule_id = "must_connect_public_underground"
    elif level == "尽连地块" and underground == "commercial_public_pr":
        rule_id = "priority_connect_public_underground"
    elif land_use in {"office_school_residential", "general_public_commercial"}:
        rule_id = "residential_or_apartment_surface"
    else:
        rule_id = "fallback_surface"

    return next(rule for rule in RULES["recommendationRules"] if rule["id"] == rule_id)


def type_by_id(type_id: str) -> dict:
    return next(item for item in RULES["connectionTypes"] if item["id"] == type_id)


def gap_detail(factor_id: str, station_context: dict | None = None) -> tuple[str, str]:
    station_context = station_context or {}
    operations = station_context.get("operations") or {}
    amenities = station_context.get("amenities") or {}
    if factor_id == "land_use":
        has_context = bool((operations.get("details") or []) or (amenities.get("nearby") or {}))
        if has_context:
            return (
                "未提供具体地块用地性质；运营表中的地块名称和周边配套已作为背景线索入库，但不能等同正式用地功能评分。",
                "缺正式地块用地性质；周边配套/地块名称仅作预评估线索"
            )
        return (
            "未提供具体地块用地性质，当前按保守 1 分临时试算。",
            "缺正式地块用地性质"
        )
    if factor_id == "development_intensity":
        return (
            "未提供建筑面积、用地面积、容积率或开发强度分档，当前按保守 1 分临时试算。",
            "缺建筑面积/容积率/开发强度正式指标"
        )
    if factor_id == "underground_space":
        forms = "、".join(operations.get("connectionForms") or [])
        if forms:
            return (
                f"运营表已识别联通形式：{forms}；但未提供地块地下空间功能、开放属性和接口标高，暂不作为地下空间因子正式赋值。",
                "有联通形式线索，缺地块地下空间正式条件"
            )
        return (
            "未提供地块地下空间功能、公共开放属性和接口条件，当前按保守 1 分临时试算。",
            "缺地块地下空间正式条件"
        )
    return (
        f"{factor_id}缺少可计算输入，当前按 1 分临时试算。",
        "系统保守预设，需补齐复核"
    )


def make_factor_value(
    factor: dict,
    value: str | None,
    missing: list[dict],
    source: str | None = None,
    station_context: dict | None = None
) -> dict:
    options = FACTOR_LOOKUP[factor["id"]]
    if value in options:
        option = options[value]
        return {
            "factorId": factor["id"],
            "name": factor["name"],
            "value": value,
            "label": option["label"],
            "score": option["score"],
            "weight": factor["weight"],
            "weightedScore": round(option["score"] * factor["weight"], 4),
            "assumed": False,
            "source": source or "现有输入/预设数据"
        }

    message, fallback_source = gap_detail(factor["id"], station_context)
    fallback = {
        "factorId": factor["id"],
        "name": factor["name"],
        "value": None,
        "label": "正式地块资料待补齐，暂按保守 1 分试算",
        "score": 1,
        "weight": factor["weight"],
        "weightedScore": round(factor["weight"], 4),
        "assumed": True,
        "source": fallback_source
    }
    missing.append({
        "factorId": factor["id"],
        "name": factor["name"],
        "message": message
    })
    return fallback


def evaluate_project(project: dict, research_options: dict | None = None) -> dict:
    station = project.get("station") or {}
    parcel = project.get("parcel") or {}
    preset_station = STATION_INDEX.get(station.get("name", ""))
    station_context = resolve_station_context(station)
    missing: list[dict] = []

    values = {
        "land_use": parcel.get("landUse"),
        "development_intensity": parcel.get("developmentIntensity") or classify_development_intensity(parcel.get("buildingArea")),
        "location_level": infer_location(station, preset_station),
        "ridership": classify_ridership(station_context.get("dailyInbound")),
        "transfer": infer_transfer(station, preset_station),
        "underground_space": parcel.get("undergroundSpace")
    }
    sources = {
        "ridership": station_context.get("dailyInboundSource") or None
    }

    dimension_results = []
    factor_results = []
    total_score = 0.0

    for dimension in FACTORS["dimensions"]:
        dimension_score = 0.0
        children = []
        for factor in dimension["factors"]:
            result = make_factor_value(
                factor,
                values.get(factor["id"]),
                missing,
                sources.get(factor["id"]),
                station_context
            )
            children.append(result)
            factor_results.append(result)
            dimension_score += result["weightedScore"]
        dimension_results.append({
            "id": dimension["id"],
            "name": dimension["name"],
            "weight": dimension["weight"],
            "score": round(dimension_score, 4),
            "factors": children
        })
        total_score += dimension_score

    score = round(total_score, 4)
    score_percent = normalized_score(score)
    grade = grade_score(score, score_percent)
    rule = pick_rule(grade["level"], parcel, missing, station, station_context)
    primary = type_by_id(rule["primary"])
    alternatives = [type_by_id(item) for item in rule["alternatives"]]
    recommendation_payload = {
        "primary": primary,
        "alternatives": alternatives,
        "rule": rule
    }

    llm_context = build_llm_review_context(
        project,
        score,
        score_percent,
        grade,
        dimension_results,
        primary,
        alternatives,
        rule,
        missing,
        station_context
    )
    research_bundle = build_model_oriented_research(
        project,
        score,
        score_percent,
        grade,
        dimension_results,
        recommendation_payload,
        missing,
        station_context,
        search_knowledge,
        research_options
    )
    model_judgement = research_bundle.get("modelJudgement") or {}
    model_rule_difference = research_bundle.get("modelRuleDifference") or {}
    model_quality = research_bundle.get("modelQuality") or {}
    diagram_brief = build_diagram_brief(project, {
        "level": grade["level"],
        "recommendation": recommendation_payload,
    }, model_judgement)
    report_modes = build_report_modes(model_judgement, model_rule_difference, model_quality)
    client_report_bundle = build_client_report(
        project,
        score,
        score_percent,
        grade,
        dimension_results,
        recommendation_payload,
        missing,
        station_context,
        research_bundle
    )
    report = build_report(
        project,
        score,
        score_percent,
        grade,
        dimension_results,
        primary,
        alternatives,
        rule,
        missing,
        station_context,
        llm_context,
        research_bundle
    )
    completeness_total = len(factor_results)
    completeness_done = sum(1 for item in factor_results if not item["assumed"])

    return {
        "project": project,
        "score": score,
        "scorePercent": score_percent,
        "scoreScale": FACTORS.get("scoreScale") or {"rawMax": max_weighted_score(), "normalizedMax": 100},
        "level": grade["level"],
        "policy": grade["policy"],
        "stationContext": station_context,
        "provisional": bool(missing),
        "dataCompleteness": {
            "done": completeness_done,
            "total": completeness_total,
            "ratio": round(completeness_done / completeness_total, 2)
        },
        "dimensions": dimension_results,
        "missing": missing,
        "recommendation": recommendation_payload,
        "researchPlan": research_bundle["researchPlan"],
        "evidencePack": research_bundle["evidencePack"],
        "modelAssessment": research_bundle["modelAssessment"],
        "modelJudgement": model_judgement,
        "modelRuleDifference": model_rule_difference,
        "modelQuality": model_quality,
        "diagramBrief": diagram_brief,
        "reportModes": report_modes,
        "capabilityStatus": research_bundle["capabilityStatus"],
        "clientReport": client_report_bundle["clientReport"],
        "clientReportMode": client_report_bundle["clientReportMode"],
        "llmReviewContext": llm_context,
        "report": report
    }


def join_cn(items: list[str], fallback: str = "暂无结构化记录") -> str:
    values = [str(item).strip() for item in items if item and str(item).strip()]
    return "；".join(values) if values else fallback


def factor_line(factor: dict) -> str:
    mark = "临时试算" if factor.get("assumed") else "已取值"
    return (
        f"{factor.get('name')}={factor.get('label')}，{factor.get('score')}分，"
        f"权重{factor.get('weight')}，加权{factor.get('weightedScore')}，来源：{factor.get('source')}，状态：{mark}"
    )


def build_reference_basis(project: dict, station_context: dict, missing: list[dict]) -> list[dict]:
    station = project.get("station") or {}
    references = [
        {
            "source": "互联互通智能体研发工作方案0318.pptx",
            "use": "确定近期目标、评价到报告的工作流、规则优先级和试点交付边界。",
            "status": "PPT规则源，优先于早期示例报告。"
        },
        {
            "source": "评价因子赋值明细表.xlsx",
            "use": "确定功能、区位、交通、地下四类维度、因子权重和选项赋值。",
            "status": "当前评分计算的直接依据。"
        },
        {
            "source": "20260507反馈",
            "use": "确认评分口径以Excel/PPT为准，早期样例报告仅作章节和表达参考。",
            "status": "验收口径确认。"
        },
        {
            "source": "车站TOD级别划分整理.xlsx",
            "use": f"提供{station.get('name') or '当前站点'}的TOD级别、线路和区位能级。",
            "status": "站点级预评估依据。"
        },
        {
            "source": "每站每月日均进站.xlsx",
            "use": "提供站点日均进站客流，用于交通属性中的出行客流因子。",
            "status": station_context.get("dailyInboundSource") or "未匹配客流时列为待补齐。"
        },
        {
            "source": "苏州地铁运营站点出入口、联通接口情况梳理统计表20260128.xlsx",
            "use": "提供出入口、接口、联通形式、开发进度和推进问题线索。",
            "status": "用于站点上下文和推荐规则分层，不替代正式地块资料。"
        },
        {
            "source": "运营公司车站出入口开放及周边配套统计表202604.xlsx",
            "use": "提供出入口开放、运营管理、周边学校/商业/住宅/公共设施等配套线索。",
            "status": "用于周边服务对象判断和资料补齐提示。"
        },
        {
            "source": "苏州轨道交通站点周边地下公共联通空间设计细则-结题评审修改稿.docx",
            "use": "作为地下联通空间、接口、无障碍、消防、运营界面等设计控制要求的知识库来源。",
            "status": "设计建议引用源，具体条款需在深化阶段人工复核。"
        }
    ]
    if missing:
        references.append({
            "source": "项目提资资料",
            "use": "补齐用地性质、建筑面积/容积率、地下空间功能、接口标高、权属边界和专项约束。",
            "status": f"当前仍有{len(missing)}项正式评分或项目资料待补齐。"
        })
    return references


def build_dimension_judgement(dimensions: list[dict]) -> list[dict]:
    judgements = []
    for dimension in dimensions:
        factors = dimension.get("factors") or []
        assumed = [factor.get("name") for factor in factors if factor.get("assumed")]
        score = float(dimension.get("score") or 0)
        if assumed:
            interpretation = f"{dimension.get('name')}含{join_cn(assumed, '')}临时试算，结论需随资料补齐复核。"
        elif score >= float(dimension.get("weight") or 0) * 3:
            interpretation = f"{dimension.get('name')}支撑度较强，对联通必要性形成正向贡献。"
        else:
            interpretation = f"{dimension.get('name')}贡献中等或偏弱，需结合工程条件和实施成本综合判断。"
        judgements.append({
            "id": dimension.get("id"),
            "name": dimension.get("name"),
            "weight": dimension.get("weight"),
            "score": dimension.get("score"),
            "factorEvidence": [factor_line(factor) for factor in factors],
            "interpretation": interpretation
        })
    return judgements


def build_scheme_comparison(primary: dict, alternatives: list[dict], rule: dict) -> list[dict]:
    schemes = [("推荐方案", primary)] + [(f"备选方案{index}", item) for index, item in enumerate(alternatives, 1)]
    comparison = []
    for role, scheme in schemes:
        comparison.append({
            "role": role,
            "name": scheme.get("name"),
            "category": scheme.get("category"),
            "reason": rule.get("reason") if role == "推荐方案" else "作为备选，用于接口、投资、交通组织或风貌约束变化时复核。",
            "bestFor": scheme.get("bestFor") or [],
            "avoidWhen": scheme.get("avoidWhen") or [],
            "technicalParameters": scheme.get("parameters") or []
        })
    return comparison


def build_diagram_brief(project: dict, result_facts: dict, judgement: dict) -> dict:
    station = (project.get("station") or {}).get("name") or "轨道站点"
    parcel = (project.get("parcel") or {}).get("name") or project.get("name") or "目标地块"
    scheme = (
        judgement.get("recommendedType")
        or ((result_facts.get("recommendation") or {}).get("primary") or {}).get("name")
        or "推荐联通方式"
    )
    return {
        "diagramId": "D-001",
        "diagramType": "recommended_connection_path",
        "title": "推荐联通路径示意",
        "summary": f"表达{station}与{parcel}之间的{scheme}关系，并标注主要风险与待复核界面。",
        "nodes": [
            {"id": "station", "label": station, "type": "station", "x": 110, "y": 160, "level": "B1"},
            {"id": "parcel", "label": parcel, "type": "parcel", "x": 430, "y": 160, "level": "B1"},
        ],
        "edges": [
            {"id": "path-1", "from": "station", "to": "parcel", "type": "recommended", "label": scheme}
        ],
        "annotations": [
            {"id": "risk-1", "text": "接口标高、消防分区、产权界面需复核", "anchorTo": "path-1"}
        ],
        "exports": ["svg", "png", "docx"],
    }


def build_report_modes(judgement: dict, difference: dict, quality: dict) -> list[dict]:
    return [
        {
            "id": "client_formal",
            "name": "客户正式版",
            "tone": "正式、稳健、可审阅",
            "focus": "结论、依据、风险、补资和实施建议",
            "sectionDensity": "medium",
            "qualityLabels": quality.get("labels", []),
        },
        {
            "id": "expert_appendix",
            "name": "专家附录版",
            "tone": "证据完整、规则透明、便于复核",
            "focus": "规则基线、模型差异、证据引用和复核问题",
            "sectionDensity": "high",
            "qualityLabels": difference.get("reviewLabels", []),
        },
        {
            "id": "leadership_brief",
            "name": "领导汇报版",
            "tone": "一页一结论、一页一风险、一页一行动",
            "focus": "模型结论、推荐方案、关键风险和下一步动作",
            "sectionDensity": "low",
            "qualityLabels": ["可汇报", *quality.get("labels", [])[:2]],
        },
    ]


def build_risk_points(project: dict, missing: list[dict], station_context: dict) -> list[dict]:
    operations = station_context.get("operations") or {}
    amenities = station_context.get("amenities") or {}
    risks = [
        {
            "topic": "正式地块资料",
            "point": "用地性质、建筑面积/容积率、地下空间功能缺失时，评分和推荐只能作为站点级预评估。",
            "action": "补齐控规图则、规划条件、地块指标和地下空间说明后复核。"
        },
        {
            "topic": "接口工程条件",
            "point": "接口位置、标高、非付费区关系、结构边界和防水节点决定推荐方式能否落地。",
            "action": "深化阶段需核对车站接口条件、地块地下室平面和竖向标高。"
        },
        {
            "topic": "消防与疏散",
            "point": "地下或封闭连通空间涉及防火分区、排烟、疏散距离、运营管理边界。",
            "action": "方案比选时同步校核消防专项和运营分界。"
        },
        {
            "topic": "无障碍与导向",
            "point": "步行接驳应保证连续无障碍、夜间照明、导向识别和客流组织清晰。",
            "action": "近期实施方案应先落实连续路径和标识系统。"
        },
        {
            "topic": "权属与施工界面",
            "point": "红线、产权、管线迁改、施工组织和费用分摊可能改变方案优先级。",
            "action": "把权属边界、施工时序和投资责任纳入补齐清单。"
        }
    ]
    if operations.get("issueCount"):
        risks.append({
            "topic": "既有推进问题",
            "point": f"运营资料记录问题{operations.get('issueCount')}条，可能影响接口开放或协议签署。",
            "action": "优先复核问题记录、接口费、联通协议和建设状态。"
        })
    if amenities and amenities.get("openExitCount") != amenities.get("exitRows"):
        risks.append({
            "topic": "出入口开放状态",
            "point": f"运营配套资料显示开放{amenities.get('openExitCount')}/{amenities.get('exitRows')}个。",
            "action": "核实现状可达性和未开放口对接条件。"
        })
    if missing:
        risks.append({
            "topic": "补齐后结论漂移",
            "point": f"当前{len(missing)}项缺口按保守值试算，补齐后评分等级和推荐方式可能变化。",
            "action": "保留本次预评估版本，补齐后生成复核版并对比分值变化。"
        })
    return risks


def build_llm_review_context(
    project: dict,
    score: float,
    score_percent: float,
    grade: dict,
    dimensions: list[dict],
    primary: dict,
    alternatives: list[dict],
    rule: dict,
    missing: list[dict],
    station_context: dict | None = None
) -> dict:
    station_context = station_context or {}
    project_name = project.get("name") or "待命名项目"
    references = build_reference_basis(project, station_context, missing)
    dimension_judgement = build_dimension_judgement(dimensions)
    scheme_comparison = build_scheme_comparison(primary, alternatives, rule)
    risk_points = build_risk_points(project, missing, station_context)
    return {
        "purpose": "供模型在规则基线、项目事实和证据材料基础上做最终综合研判，可在说明理由和复核标签后覆盖规则结论。",
        "guardrails": [
            "可以覆盖规则等级或推荐方式，但必须保留规则基线、覆盖理由、证据引用和人工复核标签。",
            "正式地块资料缺失时必须标注预评估性质，不得把站点周边线索当作地块正式输入。",
            "引用资料应说明来源和用途；无法确认的工程条件应进入补齐项或风险项。",
            "输出需区分近期可实施措施、中期深化设计和远期一体化预留。"
        ],
        "projectFacts": {
            "name": project_name,
            "station": (project.get("station") or {}).get("name"),
            "score": score,
            "scorePercent": round(score_percent, 2),
            "level": grade.get("level"),
            "policy": grade.get("policy"),
            "primaryRecommendation": primary.get("name"),
            "ruleReason": rule.get("reason"),
            "missingItems": [item.get("message") for item in missing]
        },
        "referenceBasis": references,
        "dimensionJudgement": dimension_judgement,
        "schemeComparison": scheme_comparison,
        "riskAndReviewPoints": risk_points,
        "suggestedOutputSchema": [
            "综合判断结论",
            "主要依据与引用来源",
            "推荐方案适用性说明",
            "备选方案触发条件",
            "风险与需补齐资料",
            "近期/中期/远期实施建议",
            "人工复核问题清单"
        ]
    }


def build_report(
    project: dict,
    score: float,
    score_percent: float,
    grade: dict,
    dimensions: list[dict],
    primary: dict,
    alternatives: list[dict],
    rule: dict,
    missing: list[dict],
    station_context: dict | None = None,
    llm_context: dict | None = None,
    research_bundle: dict | None = None
) -> list[dict]:
    station = project.get("station") or {}
    parcel = project.get("parcel") or {}
    station_context = station_context or {}
    operations = station_context.get("operations") or {}
    ridership = station_context.get("ridership") or {}
    amenities = station_context.get("amenities") or {}
    station_name = station.get("name") or "待补齐车站"
    project_name = project.get("name") or "待命名项目"
    missing_text = "；".join(item["message"] for item in missing) if missing else "当前核心评价字段完整"
    dimension_text = "；".join(f"{item['name']} {item['score']:.4f} 分" for item in dimensions)
    alt_text = "、".join(item["name"] for item in alternatives)
    params = "；".join(primary.get("parameters", [])[:4])
    raw_max = max_weighted_score()
    district = station.get("district") or "、".join((operations.get("districts") or [])[:2]) or "待补齐区划"
    site_area = parcel.get("siteArea") or "待补齐"
    building_area = parcel.get("buildingArea") or "待补齐"
    location = parcel.get("location") or parcel.get("quadrant") or "待补齐位置"
    land_use_text = parcel.get("functionalFormat") or parcel.get("landUseText") or "待补齐"
    planning = parcel.get("planningIndicators") or "待补齐规划指标"
    daily = station_context.get("dailyInbound")
    daily_text = (
        f"{daily:.0f} 人次/日（{station_context.get('dailyInboundSource')}）"
        if isinstance(daily, (int, float)) else "待补齐"
    )
    operation_text = (
        f"站点侧已识别出入口 {operations.get('exitCount', 0)} 个，"
        f"联通形式 {('、'.join(operations.get('connectionForms') or []) or '暂无结构化记录')}，"
        f"问题记录 {operations.get('issueCount', 0)} 条。"
        if operations else "出入口与接口信息待补齐。"
    )
    amenities_text = (
        f"运营公司202604资料显示，站点出入口开放 {amenities.get('openExitCount', 0)}/{amenities.get('exitRows', 0)} 个，"
        f"运营管理 {amenities.get('managedExitCount', 0)} 个；周边配套包括{compact_nearby_places(amenities) or '待补齐'}。"
        if amenities else "运营侧出入口开放和周边配套资料待补齐。"
    )
    ridership_note = (
        f"客流记录覆盖至 {ridership.get('latestMonth')}，平均日均进站 {ridership.get('averageDailyInbound')} 人次/日。"
        if ridership else "客流未匹配到站点预设记录。"
    )
    llm_context = llm_context or {}
    research_bundle = research_bundle or {}
    research_plan = research_bundle.get("researchPlan") or {}
    evidence_pack = research_bundle.get("evidencePack") or {}
    model_assessment = research_bundle.get("modelAssessment") or {}
    model_judgement = research_bundle.get("modelJudgement") or {}
    model_rule_difference = research_bundle.get("modelRuleDifference") or {}
    capability_status = research_bundle.get("capabilityStatus") or {}
    reference_text = "；".join(
        f"{item.get('source')}（{item.get('use')}）"
        for item in (llm_context.get("referenceBasis") or [])[:8]
    )
    dimension_judgement_text = "；".join(
        f"{item.get('name')}：{item.get('interpretation')}"
        for item in (llm_context.get("dimensionJudgement") or [])
    )
    comparison_text = "；".join(
        f"{item.get('role')}{item.get('name')}，适用：{join_cn(item.get('bestFor') or [], '待结合项目条件判断')}，回避：{join_cn(item.get('avoidWhen') or [], '暂无')}"
        for item in (llm_context.get("schemeComparison") or [])
    )
    risk_text = "；".join(
        f"{item.get('topic')}：{item.get('point')}建议：{item.get('action')}"
        for item in (llm_context.get("riskAndReviewPoints") or [])[:8]
    )
    llm_schema_text = "、".join(llm_context.get("suggestedOutputSchema") or [])
    capability_text = (
        f"本地缓存{(capability_status.get('localCache') or {}).get('mode', '未知')}，"
        f"独立搜索{(capability_status.get('independentSearch') or {}).get('mode', '未知')}，"
        f"模型{(capability_status.get('llm') or {}).get('mode', '未知')}，"
        f"模型联网{(capability_status.get('modelWebSearch') or {}).get('mode', '未知')}。"
    )
    research_questions_text = "；".join(
        item.get("question", "")
        for item in (research_plan.get("questions") or [])[:6]
        if item.get("question")
    )
    evidence_text = "；".join(
        f"{item.get('title')}（{item.get('source')}，置信度{item.get('confidence')}，{'缓存' if item.get('cached') else '实时'}）"
        for item in (evidence_pack.get("items") or [])[:10]
    )
    model_dimension_text = "；".join(
        f"{item.get('name')}：{item.get('judgement') or item.get('summary') or item.get('focus')}（置信度{item.get('confidence', '待确认')}）"
        for item in (model_assessment.get("dynamicDimensions") or [])[:8]
    )
    model_review_text = "；".join(model_assessment.get("reviewQuestions") or [])
    model_uncertainty_text = "；".join(model_assessment.get("uncertainties") or [])

    return [
        {
            "title": "项目概况与资料边界",
            "content": (
                f"{project_name}对应站点为{station_name}，所属区划为{district}，地块位置为{location}，"
                f"圈层为{parcel.get('distanceBand') or '待补齐'}。地块业态为{land_use_text}，"
                f"用地面积约{site_area}平方米，建筑面积约{building_area}平方米，规划指标为{planning}。"
                f"邻近出入口为{station.get('nearbyExit') or '待补齐'}，日均客流为{daily_text}。{operation_text}{amenities_text}"
                "本段仅描述已入库事实和待补齐资料边界，不把周边配套线索直接等同为地块正式评分输入。"
            )
        },
        {
            "title": "评价口径与引用依据",
            "content": (
                f"本次评价以PPT工作方案、评价因子赋值表和20260507反馈为主控口径，引用依据包括：{reference_text}。"
                "其中评分因子、权重和等级阈值由结构化规则引擎计算；设计细则、运营资料和周边配套作为解释、比选、风险识别和补齐项提示依据。"
            )
        },
        {
            "title": "站点能级、客流与接口研判",
            "content": (
                f"{station_name}的区位能级为{station.get('todLevel') or station.get('locationLevel') or '待补齐'}，"
                f"线路/换乘信息为{station.get('line') or '待补齐'}，日均客流为{daily_text}。{ridership_note}"
                f"{operation_text}这些信息用于判断站点侧联通需求、接口成熟度和近期实施抓手；若后续接入客流预测、客流方向、分时段拥挤度，可进一步细化为路径级方案建议。"
            )
        },
        {
            "title": "周边功能与服务对象研判",
            "content": (
                f"运营公司202604资料和站点周边索引显示，周边配套包括{compact_nearby_places(amenities) or '待补齐'}。"
                "这些内容可辅助判断服务对象、全天候需求、导向标识和步行系统连续性，但仍需控规、规划条件和地块方案确认具体用地功能、开发强度和地下空间条件。"
            )
        },
        {
            "title": "多维度评分拆解",
            "content": (
                f"本次按功能、区位、交通、地下四个维度进行参数化试算，维度得分为：{dimension_text}。"
                f"维度解释为：{dimension_judgement_text}。"
                "对临时试算因子，系统保留1分保守值并在待补齐项中说明原因，避免用推断数据替代正式资料。"
            )
        },
        {
            "title": "联通必要性评估结论",
            "content": (
                f"原始加权分为 {score:.4f}/{raw_max:.4f}，折算百分制为 {score_percent:.2f} 分，"
                f"按 20260507 反馈的 80/60 阈值判定为{grade['level']}。{grade['policy']}{ridership_note}"
                "该结论表示当前资料条件下的必要性等级；若后续补齐地块规模、地下空间或接口条件，系统应重新计算并输出差异说明。"
            )
        },
        {
            "title": "推荐方案与方案比选",
            "content": (
                f"系统推荐采用{primary['name']}，推荐理由为：{rule['reason']}。"
                f"备选方案包括{alt_text}。比选框架为：{comparison_text}。"
                "后续可按接口可达性、投资强度、施工影响、全天候服务、运营管理和城市风貌六类指标做专家复核。"
            )
        },
        {
            "title": "方案设计核心技术要求",
            "content": (
                f"{primary['name']}属于{primary['category']}方式。核心技术要求包括：{params}。"
                "深化设计还应明确通行净宽、净高、坡道/电梯、导向标识、照明、排水、防滑、运营管理界面和应急疏散组织。"
            )
        },
        {
            "title": "风险约束与专项复核",
            "content": (
                f"重点风险和复核建议包括：{risk_text}。"
                "上述风险不直接否定推荐方案，但会影响方案层级、接口形式、建设时序和投资责任划分。"
            )
        },
        {
            "title": "智能研究计划与能力状态",
            "content": (
                f"系统按“规则保底、实时研究、本地知识兜底、模型增强”的策略生成研究计划，当前能力状态为：{capability_text}"
                f"匹配标杆案例为{(research_plan.get('benchmarkCase') or {}).get('label', '通用站点互联互通预评估')}。"
                f"本轮研究问题包括：{research_questions_text}。"
                "模型可用时以模型综合研判为主，本地资料作为证据弹药；未连接模型时需用户确认后才使用离线兜底。"
            )
        },
        {
            "title": "证据包与来源置信度",
            "content": (
                f"本轮证据包模式为{evidence_pack.get('mode', 'offline_cache')}，"
                f"共收集{(evidence_pack.get('summary') or {}).get('total', 0)}条证据，"
                f"其中缓存{(evidence_pack.get('summary') or {}).get('cached', 0)}条、实时{(evidence_pack.get('summary') or {}).get('live', 0)}条。"
                f"代表性证据包括：{evidence_text}。"
                "证据优先级为项目正式输入/规则源、本地标杆缓存、独立搜索、模型联网低置信补充；未被印证的外部信息只作为复核线索。"
            )
        },
        {
            "title": "模型导向多维研判",
            "content": (
                f"{model_assessment.get('summary') or '当前采用本地兜底研判。'}"
                f"动态维度研判包括：{model_dimension_text}。"
                f"不确定性包括：{model_uncertainty_text or '暂无额外不确定性'}。"
                f"人工复核问题包括：{model_review_text}。"
                "上述模型导向研判以模型为最终综合判断，规则评分作为基线和差异检查坐标保留。"
            )
        },
        {
            "title": "模型主导综合研判",
            "content": (
                f"模型主结论为：{model_judgement.get('level', grade['level'])}，推荐方式为"
                f"{model_judgement.get('recommendedType', primary.get('name'))}。"
                f"置信度约为{model_judgement.get('confidence', 0):.2f}。"
                f"判断理由：{model_judgement.get('reason', '模型基于规则基线、站点上下文和资料证据形成综合判断。')}"
                f"差异状态：{model_rule_difference.get('status', 'aligned')}；"
                f"复核标签：{join_cn(model_rule_difference.get('reviewLabels') or [])}。"
            )
        },
        {
            "title": "资料补齐与复核路径",
            "content": (
                f"资料完整性状态：{missing_text}。建议在方案深化前补齐站点客流、接口标高、出入口位置、地下空间功能、"
                "建设时序和特殊约束条件。补齐后应重新运行评分，并输出补齐前后等级、推荐方式和关键因子变化。"
            )
        },
        {
            "title": "实施时序与协同建议",
            "content": (
                "近期建议先形成可落地的步行接驳、导向标识、接口预留和资料补齐清单；中期结合地块方案、地下空间和接口标高做工程可行性比选；"
                "远期在片区更新或地块开发同步阶段推动站城一体化连通、运营管理界面和投资责任机制。"
                "责任协同建议覆盖轨道运营、规划设计、地块开发主体、道路/市政/管线单位和消防审查单位。"
            )
        },
        {
            "title": "LLM综合判断框架",
            "content": (
                "后续接入其他LLM时，建议把本次后端结果作为规则基线、证据包和差异检查输入，LLM负责最终综合研判、方案复核和专家式建议。"
                f"输入应包括项目事实、引用依据、四维评分拆解、推荐与备选方案、风险复核点和待补齐资料。建议输出结构为：{llm_schema_text}。"
                "模型可以在明确说明理由、证据和复核标签的前提下覆盖规则等级或推荐方案；规则评分作为基线和差异检查坐标保留。"
                "模型不得把缺失资料补成确定事实；所有新增判断应标注依据、置信度和需人工确认的问题。"
            )
        }
    ]


def build_markdown_report(result: dict) -> str:
    project = result.get("project") or {}
    title = project.get("name") or "互联互通评估报告"
    primary = ((result.get("recommendation") or {}).get("primary") or {}).get("name") or "待推荐"
    raw_max = ((result.get("scoreScale") or {}).get("rawMax") or max_weighted_score())
    lines = [
        f"# {title}",
        "",
        f"- 地块编号：{project.get('projectCode') or '待补齐'}",
        f"- 综合评分（百分制）：{float(result.get('scorePercent') or 0):.2f}分",
        f"- 原始加权分：{float(result.get('score') or 0):.4f} / {float(raw_max):.4f}",
        f"- 联通等级：{result.get('level') or '待判定'}",
        f"- 推荐方式：{primary}",
        f"- 数据完整度：{(result.get('dataCompleteness') or {}).get('done', 0)}/{(result.get('dataCompleteness') or {}).get('total', 0)}",
        ""
    ]
    lines.extend([
        "## 评分总览",
        "",
        "| 项目 | 结果 |",
        "| --- | --- |",
        f"| 综合评分（百分制） | {float(result.get('scorePercent') or 0):.2f}分 |",
        f"| 原始加权分 | {float(result.get('score') or 0):.4f} / {float(raw_max):.4f} |",
        f"| 联通等级 | {result.get('level') or '待判定'} |",
        f"| 推荐方式 | {primary} |",
        f"| 数据完整度 | {(result.get('dataCompleteness') or {}).get('done', 0)}/{(result.get('dataCompleteness') or {}).get('total', 0)} |",
        "",
        "| 评分维度 | 加权得分 | 主要因子取值 | 来源依据 | 说明 |",
        "| --- | ---: | --- | --- | --- |",
    ])
    for row in score_overview_rows(result):
        lines.append(
            f"| {row['dimension']} | {row['score']:.4f} | "
            f"{row['keyFactors'].replace('|', '/')} | "
            f"{row['sourceBasis'].replace('|', '/')} | {row['note']} |"
        )
    lines.append("")
    client_sections = result.get("clientReport") or result.get("report") or []
    for index, section in enumerate(client_sections, 1):
        lines.extend([f"## {index}. {section.get('title')}", "", section.get("content") or "", ""])
    missing = result.get("missing") or []
    if missing:
        lines.extend(["## 资料补齐与复核事项", ""])
        for item in missing:
            lines.append(f"- {item.get('name')}：{item.get('message')}")
        lines.append("")
    return "\n".join(lines)


def save_project_record(project: dict, research_options: dict | None = None) -> dict:
    research_options = research_options or {}
    data = load_projects()
    now = utc_now()
    record_id = project.get("id") or uuid.uuid4().hex[:12]
    existing = next((item for item in data.get("projects", []) if item.get("id") == record_id), None)
    result = (existing or {}).get("result")
    if not research_options.get("skipEvaluation"):
        result = evaluate_project(project, research_options)
    owner = project.get("owner") or (existing or {}).get("owner") or current_owner_metadata("project")
    project = {**project, "owner": owner}
    record = {
        "id": record_id,
        "project": project,
        "result": result,
        "owner": owner,
        "createdAt": (existing or {}).get("createdAt") or now,
        "updatedAt": now
    }
    projects = [item for item in data.get("projects", []) if item.get("id") != record_id]
    projects.append(record)
    data["projects"] = projects
    save_projects(data)
    return record


def migrate_project_owner(record_id: str, target_owner: dict | None = None) -> dict:
    data = load_projects()
    record = next((item for item in data.get("projects", []) if item.get("id") == record_id), None)
    if not record:
        raise ValueError("Project not found for ownership migration")
    owner = migrate_owner_metadata(record, target_owner)
    record["owner"] = owner
    record.setdefault("ownerMigrations", [])
    record["ownerMigrations"].append(owner["migration"])
    project = record.get("project") or {}
    project["owner"] = owner
    record["project"] = project
    record["updatedAt"] = utc_now()
    save_projects(data)
    return record


def find_project_record(record_id: str, refresh: bool = False) -> dict | None:
    record = next((item for item in load_projects().get("projects", []) if item.get("id") == record_id), None)
    if record and refresh:
        record["result"] = evaluate_project(record.get("project") or {})
    return record


def delete_project_record(record_id: str) -> bool:
    data = load_projects()
    before = len(data.get("projects", []))
    data["projects"] = [item for item in data.get("projects", []) if item.get("id") != record_id]
    if len(data["projects"]) == before:
        return False
    save_projects(data)
    return True


def resolve_project_payload(payload: dict) -> tuple[dict, dict]:
    if isinstance(payload.get("project"), dict):
        return payload["project"], payload.get("researchOptions") or {}
    return payload, payload.get("researchOptions") or {}


def resolve_export_result(payload: dict) -> dict:
    project = payload.get("project")
    project_id = payload.get("projectId")
    legacy_result = payload.get("result")
    research_options = payload.get("researchOptions") or {}

    if project_id:
        record = find_project_record(str(project_id))
        if not record:
            raise ValueError("Project not found for export")
        project = record.get("project") or {}
    elif not isinstance(project, dict) and isinstance(legacy_result, dict):
        project = legacy_result.get("project")
    elif not isinstance(project, dict):
        project = payload

    if not isinstance(project, dict):
        raise ValueError("Export requires a project payload")
    return evaluate_project(project, research_options)


def require_export_artifact(path: Path | None, artifact_class: str) -> Path:
    if not path:
        raise RuntimeError(f"Export artifact failed: {artifact_class} was not generated")
    if not path.exists() or path.stat().st_size <= 0:
        raise RuntimeError(f"Export artifact failed: {artifact_class} is missing or empty")
    return path


def export_pdf_enabled() -> bool:
    return truthy_env("INTERCONNECT_EXPORT_PDF") or truthy_env("EXPORT_PDF")


def export_report_file(result: dict) -> dict:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    project = result.get("project") or {}
    slug = safe_slug(project.get("projectCode") or project.get("name"), "interconnect-report")
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S-%f")

    files = []
    snapshot_file = EXPORT_DIR / f"{slug}-{stamp}-evaluation-snapshot.json"
    snapshot_file.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    snapshot_file = require_export_artifact(snapshot_file, "evaluation-snapshot.json")
    report_docx = create_docx_report(result, slug, stamp)
    report_docx = require_export_artifact(report_docx, "formal-report.docx")
    score_docx = create_score_detail_docx(result, slug, stamp)
    files.append(file_info(report_docx))
    files.append(file_info(snapshot_file))
    for docx_file in [report_docx, score_docx]:
        if not docx_file:
            continue
        if docx_file != report_docx:
            files.append(file_info(docx_file))
        if export_pdf_enabled():
            pdf_file = create_pdf_from_docx(docx_file)
            if pdf_file:
                files.append(file_info(pdf_file))

    export = file_info(report_docx)
    export["files"] = files
    export["snapshot"] = file_info(snapshot_file)
    export["owner"] = project.get("owner") or current_owner_metadata("export")
    return export


def csv_escape(value: str) -> str:
    if any(char in value for char in [",", "\"", "\n"]):
        return "\"" + value.replace("\"", "\"\"") + "\""
    return value


def file_info(path: Path) -> dict:
    relative_path = path.relative_to(ROOT).as_posix()
    content_type = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
    return {
        "filename": path.name,
        "path": str(path),
        "relativePath": relative_path,
        "downloadUrl": f"/{relative_path}",
        "contentType": content_type,
        "size": path.stat().st_size
    }


def schematic_png_file_info(path: Path) -> dict:
    target = path.resolve()
    export_root = SCHEMATIC_EXPORT_DIR.resolve()
    target.relative_to(export_root)
    relative_path = target.relative_to(FRONTEND_DIR).as_posix()
    return {
        "filename": target.name,
        "path": str(target),
        "relativePath": relative_path,
        "downloadUrl": f"/{relative_path}",
        "contentType": mimetypes.guess_type(str(target))[0] or "image/png",
        "size": target.stat().st_size,
    }


def schematic_png_export_response(result: dict) -> dict:
    enriched = dict(result)
    output_path = enriched.get("outputPath")
    if enriched.get("ok") and output_path:
        export = schematic_png_file_info(Path(output_path))
        enriched["export"] = export
        enriched["filename"] = export["filename"]
        enriched["relativePath"] = export["relativePath"]
        enriched["downloadUrl"] = export["downloadUrl"]
        enriched["contentType"] = export["contentType"]
        enriched["size"] = export["size"]
    return enriched


def export_kind(path: Path) -> str:
    name = path.name.lower()
    if name.endswith("-score-detail.docx"):
        return "打分明细 Word"
    if name.endswith("-score-detail.pdf"):
        return "打分明细 PDF"
    if name.endswith("-formal-report.docx"):
        return "正式报告 Word"
    if name.endswith("-formal-report.pdf"):
        return "正式报告 PDF"
    if name.endswith(".pdf"):
        return "PDF 报告"
    if name.endswith(".docx"):
        return "Word 报告"
    if name.endswith(".md"):
        return "Markdown 报告"
    if path.parent.name == "generated-images" and name.endswith(".metadata.json"):
        return "生成图元数据"
    if path.parent.name == "generated-images":
        return "生成图"
    if name.endswith(".json"):
        return "评估快照"
    if name.endswith("-score-detail.csv"):
        return "评分明细"
    if name.endswith("-missing.csv"):
        return "待补齐清单"
    if name.endswith(".csv"):
        return "CSV 数据"
    return "导出文件"


def list_export_files(limit: int = 40) -> list[dict]:
    if not EXPORT_DIR.exists():
        return []
    candidate_paths = [path for path in EXPORT_DIR.iterdir()]
    generated_dir = EXPORT_DIR / "generated-images"
    if generated_dir.exists():
        candidate_paths.extend(path for path in generated_dir.iterdir())
    files = [
        path for path in candidate_paths
        if (
            path.is_file()
            and path.suffix.lower() in {".docx", ".pdf", ".json", ".png", ".jpg", ".jpeg", ".webp", ".svg"}
            and (
                path.name.lower().endswith("-formal-report.docx")
                or path.name.lower().endswith("-formal-report.pdf")
                or path.name.lower().endswith("-score-detail.docx")
                or path.name.lower().endswith("-score-detail.pdf")
                or path.name.lower().endswith("-evaluation-snapshot.json")
                or path.parent.name == "generated-images"
            )
        )
    ]
    files.sort(key=lambda path: path.stat().st_mtime, reverse=True)
    items = []
    for path in files[:limit]:
        info = file_info(path)
        info["kind"] = export_kind(path)
        info["updatedAt"] = datetime.fromtimestamp(path.stat().st_mtime).replace(microsecond=0).isoformat()
        items.append(info)
    return items


def manifest_file_info(path: Path) -> dict:
    stat = path.stat()
    return {
        "filename": path.name,
        "relativePath": path.relative_to(ROOT).as_posix(),
        "size": stat.st_size,
        "updatedAt": datetime.fromtimestamp(stat.st_mtime).replace(microsecond=0).isoformat()
    }


def existing_files(paths: list[Path]) -> list[Path]:
    seen: set[str] = set()
    files: list[Path] = []
    for path in paths:
        if not path.exists() or not path.is_file():
            continue
        key = str(path.resolve())
        if key in seen:
            continue
        seen.add(key)
        files.append(path)
    return files


def delivery_manifest_groups() -> list[dict]:
    doc_files = existing_files([
        ROOT / "docs" / "delivery_notes.md",
        ROOT / "docs" / "acceptance_report.md",
        ROOT / "docs" / "api_reference.md",
        ROOT / "docs" / "delivery_manifest.md",
        ROOT / "docs" / "function_and_gap_inventory.md",
        ROOT / "docs" / "knowledge_database.md",
        ROOT / "docs" / "feedback_20260507_analysis.md",
        ROOT / "docs" / "llm_integration_contract.md",
        ROOT / "docs" / "pilot_input_template.xlsx",
    ])
    runtime_files = existing_files([
        ROOT / "README.md",
        ROOT / "run.ps1",
        ROOT / "backend" / "server.py",
        ROOT / "backend" / "research_agent.py",
        ROOT / "verify_system.cjs",
        ROOT / "tools" / "seed_station_projects.py",
        ROOT / "tools" / "verify_station_precheck.cjs",
        ROOT / "tools" / "verify_report_richness.cjs",
        ROOT / "tools" / "verify_delivery_package.cjs",
        ROOT / "tools" / "verify_model_oriented_research.py",
        ROOT / "tools" / "verify_client_report.py",
    ])
    screenshot_files = existing_files([
        ROOT / "interconnect_dashboard.png",
        ROOT / "interconnect_assessment.png",
        ROOT / "interconnect_knowledge.png",
        ROOT / "interconnect_reporting.png",
        ROOT / "interconnect_agent_screenshot.png",
    ])
    mockup_files = existing_files([
        ROOT / "docs" / "mockups" / "dashboard-concept-20260508.png",
        ROOT / "docs" / "mockups" / "dashboard-ideal.png",
        ROOT / "docs" / "mockups" / "assessment-ideal.png",
        ROOT / "docs" / "mockups" / "knowledge-ideal.png",
        ROOT / "docs" / "mockups" / "report-completion-ideal.png",
    ])
    export_files = [ROOT / item["relativePath"] for item in list_export_files(20)]
    groups = [
        ("runtime", "运行与验收入口", runtime_files),
        ("docs", "交付文档", doc_files),
        ("screenshots", "验收截图", screenshot_files),
        ("mockups", "理想页面视觉稿", mockup_files),
        ("exports", "最近导出成果", existing_files(export_files)),
    ]
    return [
        {
            "id": group_id,
            "title": title,
            "files": [manifest_file_info(path) for path in files],
        }
        for group_id, title, files in groups
    ]


def build_delivery_manifest_payload() -> dict:
    groups = delivery_manifest_groups()
    files = [item for group in groups for item in group["files"]]
    total_size = sum(item["size"] for item in files)
    stamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    return {
        "generatedAt": utc_now(),
        "totalFiles": len(files),
        "totalSize": total_size,
        "groups": groups,
        "files": files,
        "package": {
            "filename": f"interconnect-agent-delivery-{stamp}.zip",
            "downloadUrl": "/api/delivery/package",
        },
    }


def build_delivery_package() -> tuple[str, bytes]:
    manifest = build_delivery_manifest_payload()
    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr("delivery_manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
        for item in manifest["files"]:
            path = ROOT / item["relativePath"]
            if path.exists() and path.is_file():
                archive.write(path, item["relativePath"])
    return manifest["package"]["filename"], buffer.getvalue()


def set_doc_defaults(doc) -> None:
    from docx.shared import Pt

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)


def create_score_detail_docx(result: dict, slug: str, stamp: str) -> Path | None:
    try:
        from docx import Document
    except Exception:
        return None

    target = EXPORT_DIR / f"{slug}-{stamp}-score-detail.docx"
    project = result.get("project") or {}
    raw_max = ((result.get("scoreScale") or {}).get("rawMax") or max_weighted_score())

    doc = Document()
    set_doc_defaults(doc)
    doc.add_heading(f"{project.get('name') or '互联互通评估'}打分明细", 0)
    doc.add_paragraph(f"地块编号：{project.get('projectCode') or '待补齐'}")
    doc.add_paragraph(f"综合评分（百分制）：{float(result.get('scorePercent') or 0):.2f}分")
    doc.add_paragraph(f"原始加权分：{float(result.get('score') or 0):.4f} / {float(raw_max):.4f}")
    doc.add_paragraph(f"联通等级：{result.get('level') or '待判定'}")

    table = doc.add_table(rows=1, cols=8)
    headers = table.rows[0].cells
    for index, label in enumerate(["维度", "因子", "选项", "原始分", "权重", "加权分", "来源", "是否假定"]):
        headers[index].text = label

    for dimension in result.get("dimensions") or []:
        for factor in dimension.get("factors") or []:
            cells = table.add_row().cells
            values = [
                dimension.get("name") or "",
                factor.get("name") or "",
                factor.get("label") or "",
                str(factor.get("score") or ""),
                str(factor.get("weight") or ""),
                str(factor.get("weightedScore") or ""),
                factor.get("source") or "",
                "是" if factor.get("assumed") else "否",
            ]
            for index, value in enumerate(values):
                cells[index].text = value

    missing = result.get("missing") or []
    doc.add_heading("资料缺项", level=1)
    if missing:
        missing_table = doc.add_table(rows=1, cols=3)
        missing_headers = missing_table.rows[0].cells
        missing_headers[0].text = "因子"
        missing_headers[1].text = "名称"
        missing_headers[2].text = "说明"
        for item in missing:
            cells = missing_table.add_row().cells
            cells[0].text = item.get("factorId") or ""
            cells[1].text = item.get("name") or ""
            cells[2].text = item.get("message") or ""
    else:
        doc.add_paragraph("当前核心评分字段完整，导出前仍建议结合正式设计条件进行人工复核。")

    doc.save(target)
    return target


def create_pdf_from_docx(docx_path: Path) -> Path | None:
    pdf_path = docx_path.with_suffix(".pdf")
    if convert_docx_to_pdf_with_word(docx_path, pdf_path):
        return pdf_path
    if convert_docx_to_pdf_with_libreoffice(docx_path, pdf_path):
        return pdf_path
    if create_simple_pdf_from_docx_text(docx_path, pdf_path):
        return pdf_path
    if create_minimal_pdf_placeholder(docx_path, pdf_path):
        return pdf_path
    return None


def convert_docx_to_pdf_with_word(docx_path: Path, pdf_path: Path) -> bool:
    try:
        import pythoncom
        from win32com import client
    except Exception:
        return False

    word = None
    document = None
    try:
        pythoncom.CoInitialize()
        word = client.DispatchEx("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(docx_path.resolve()))
        document.ExportAsFixedFormat(str(pdf_path.resolve()), 17)
        return pdf_path.exists() and pdf_path.stat().st_size > 0
    except Exception:
        return False
    finally:
        try:
            if document is not None:
                document.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def convert_docx_to_pdf_with_libreoffice(docx_path: Path, pdf_path: Path) -> bool:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        return False
    try:
        subprocess.run(
            [
                executable,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf_path.parent),
                str(docx_path),
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=60,
        )
    except Exception:
        return False
    converted = docx_path.with_suffix(".pdf")
    return converted.exists() and converted.stat().st_size > 0


def create_simple_pdf_from_docx_text(docx_path: Path, pdf_path: Path) -> bool:
    try:
        import fitz
        from docx import Document
    except Exception:
        return False


def pdf_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def create_minimal_pdf_placeholder(docx_path: Path, pdf_path: Path) -> bool:
    try:
        line = pdf_escape(f"PDF fallback generated from DOCX: {docx_path.name}")
        content = f"BT /F1 12 Tf 72 760 Td ({line}) Tj ET"
        objects = [
            "1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
            "2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
            "3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
            "/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >> endobj\n",
            f"4 0 obj << /Length {len(content.encode('latin-1'))} >> stream\n{content}\nendstream endobj\n",
            "5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        ]
        data = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for obj in objects:
            offsets.append(len(data))
            data.extend(obj.encode("latin-1"))
        xref_offset = len(data)
        data.extend(f"xref\n0 {len(objects) + 1}\n".encode("latin-1"))
        data.extend(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            data.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
        data.extend(
            (
                f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\n"
                f"startxref\n{xref_offset}\n%%EOF\n"
            ).encode("latin-1")
        )
        pdf_path.write_bytes(bytes(data))
        return pdf_path.exists() and pdf_path.stat().st_size > 0
    except Exception:
        return False

    try:
        source = Document(docx_path)
        lines: list[str] = []
        for paragraph in source.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.extend(textwrap.wrap(text, width=48) or [text])
                lines.append("")
        for table in source.tables:
            for row in table.rows:
                text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
                lines.extend(textwrap.wrap(text, width=48) or [text])
            lines.append("")

        pdf = fitz.open()
        page = pdf.new_page(width=595, height=842)
        x, y = 54, 54
        for line in lines or [docx_path.stem]:
            if y > 790:
                page = pdf.new_page(width=595, height=842)
                y = 54
            page.insert_text((x, y), line, fontname="china-s", fontsize=10)
            y += 16
        pdf.save(pdf_path)
        return pdf_path.exists() and pdf_path.stat().st_size > 0
    except Exception:
        return False


def create_docx_report(result: dict, slug: str, stamp: str) -> Path | None:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        return None

    target = EXPORT_DIR / f"{slug}-{stamp}-formal-report.docx"
    project = result.get("project") or {}
    primary = ((result.get("recommendation") or {}).get("primary") or {}).get("name") or "待推荐"

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    raw_max = ((result.get("scoreScale") or {}).get("rawMax") or max_weighted_score())
    recommendation = result.get("recommendation") or {}
    primary_obj = recommendation.get("primary") or {}
    alternatives = recommendation.get("alternatives") or []

    # —— 封面元信息（对齐旧版：编制单位 / 编制日期 / 地块编号 / 对应站点）——
    doc.add_heading(project.get("name") or "互联互通评估报告", 0)
    station = project.get("station") or {}
    station_name = str(station.get("name") or "").strip()
    if station_name and not station_name.endswith("站"):
        station_name += "站"
    line = station.get("line")
    station_full = (
        f"苏州轨道交通{line}号线{station_name}" if line and station_name else (station_name or "待补齐")
    )
    report_month = "待补齐"
    if len(str(stamp)) >= 8 and str(stamp)[:8].isdigit():
        report_month = f"{str(stamp)[:4]}年{str(stamp)[4:6]}月"
    doc.add_paragraph(f"报告编制单位：{project.get('reportOrg') or '苏州市轨道交通集团有限公司'}")
    doc.add_paragraph(f"报告编制日期：{report_month}")
    doc.add_paragraph(f"地块编号：{project.get('projectCode') or '待补齐'}")
    doc.add_paragraph(f"对应站点：{station_full}")

    client_sections = result.get("clientReport") or result.get("report") or []
    missing = result.get("missing") or []

    # —— 目录 ——
    toc_titles = [str(section.get("title") or "").strip() for section in client_sections if section.get("title")]
    doc.add_heading("目录", level=1)
    for idx, title in enumerate(toc_titles, 1):
        doc.add_paragraph(f"{idx}  {title}")
    if missing:
        doc.add_paragraph(f"{len(toc_titles) + 1}  资料补齐与复核事项")
    doc.add_page_break()

    def render_score_overview() -> None:
        # 评分总览归入“联通必要性评估结论”章节之下，对齐旧版 2.x 计算过程层次。
        doc.add_heading("评分总览", level=2)
        summary_table = doc.add_table(rows=1, cols=2)
        summary_headers = summary_table.rows[0].cells
        summary_headers[0].text = "项目"
        summary_headers[1].text = "结果"
        summary_items = [
            ("综合评分（百分制）", f"{float(result.get('scorePercent') or 0):.2f}分"),
            ("原始加权分", f"{float(result.get('score') or 0):.4f} / {float(raw_max):.4f}"),
            ("联通等级", result.get("level") or "待判定"),
            ("推荐方式", primary),
            ("数据完整度", f"{(result.get('dataCompleteness') or {}).get('done', 0)}/{(result.get('dataCompleteness') or {}).get('total', 0)}"),
        ]
        for label, value in summary_items:
            row = summary_table.add_row().cells
            row[0].text = label
            row[1].text = value

        dimension_table = doc.add_table(rows=1, cols=5)
        dimension_headers = dimension_table.rows[0].cells
        dimension_headers[0].text = "评分维度"
        dimension_headers[1].text = "加权得分"
        dimension_headers[2].text = "主要因子取值"
        dimension_headers[3].text = "来源依据"
        dimension_headers[4].text = "说明"
        for item in score_overview_rows(result):
            row = dimension_table.add_row().cells
            row[0].text = item["dimension"]
            row[1].text = f"{item['score']:.4f}"
            row[2].text = item["keyFactors"]
            row[3].text = item["sourceBasis"]
            row[4].text = item["note"]

    def render_recommendation() -> None:
        # 推荐与备选方案归入“联通方式比选与推荐方案”章节之下。
        doc.add_heading("推荐方案与备选方案", level=2)
        rec_table = doc.add_table(rows=1, cols=4)
        rec_headers = rec_table.rows[0].cells
        rec_headers[0].text = "类型"
        rec_headers[1].text = "名称"
        rec_headers[2].text = "类别"
        rec_headers[3].text = "核心适用场景"
        row = rec_table.add_row().cells
        row[0].text = "推荐"
        row[1].text = primary_obj.get("name") or ""
        row[2].text = primary_obj.get("category") or ""
        row[3].text = "、".join(primary_obj.get("bestFor") or [])
        for item in alternatives:
            row = rec_table.add_row().cells
            row[0].text = "备选"
            row[1].text = item.get("name") or ""
            row[2].text = item.get("category") or ""
            row[3].text = "、".join(item.get("bestFor") or [])

    rendered_score = False
    rendered_rec = False
    for index, section in enumerate(client_sections, 1):
        title = str(section.get("title") or "").strip()
        doc.add_heading(f"{index} {title}", level=1)
        content = str(section.get("content") or "")
        blocks = [block.strip() for block in content.split("\n") if block.strip()] or [content]
        for block in blocks:
            doc.add_paragraph(block)
        if not rendered_score and ("必要性" in title or "评估结论" in title):
            render_score_overview()
            rendered_score = True
        if not rendered_rec and ("比选" in title or "推荐方案" in title):
            render_recommendation()
            rendered_rec = True
    if not rendered_score:
        render_score_overview()
    if not rendered_rec:
        render_recommendation()

    if missing:
        doc.add_heading("资料补齐与复核事项", level=1)
        table = doc.add_table(rows=1, cols=3)
        headers = table.rows[0].cells
        headers[0].text = "因子"
        headers[1].text = "名称"
        headers[2].text = "说明"
        for item in missing:
            cells = table.add_row().cells
            cells[0].text = item.get("factorId") or ""
            cells[1].text = item.get("name") or ""
            cells[2].text = item.get("message") or ""
    doc.save(target)
    return target


class Handler(BaseHTTPRequestHandler):
    server_version = "InterconnectAgent/1.0"

    def do_OPTIONS(self) -> None:
        self.send_response(HTTPStatus.NO_CONTENT)
        self._cors()
        self.end_headers()

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        route = parsed.path
        if route == "/api/health":
            self._json({"ok": True, "service": "interconnect-agent", "version": FACTORS["version"]})
            return
        if route == "/api/capabilities":
            self._json({"ok": True, "capabilities": build_platform_capability_status()})
            return
        if route == "/api/identity":
            self._json({"ok": True, **local_identity_payload()})
            return
        if route == "/api/admin/station-outlines":
            params = parse_qs(parsed.query)
            station_name = (params.get("station") or params.get("stationName") or [""])[0] or None
            self._json({"ok": True, "records": list_admin_station_outlines(station_name)})
            return
        if route == "/api/station-memory":
            params = parse_qs(parsed.query)
            station_name = (params.get("station") or params.get("stationName") or params.get("q") or [""])[0] or None
            self._json({
                "ok": True,
                "version": load_station_memory_data().get("version"),
                "records": list_station_memory_records(station_name),
            })
            return
        if route == "/api/knowledge":
            self._json({
                "ok": True,
                "catalog": KNOWLEDGE_CATALOG,
                "ruleCards": {
                    "count": len(RULE_CARDS.get("cards", [])),
                    "sample": RULE_CARDS.get("cards", [])[:8]
                },
                "stationIndex": {
                    "count": STATION_KNOWLEDGE_INDEX.get("count", 0),
                    "coverage": STATION_KNOWLEDGE_INDEX.get("coverage", {})
                },
                "sources": source_manifest_records(),
                "unparsedSources": UNPARSED_SOURCES
            })
            return
        if route == "/api/knowledge/search":
            params = parse_qs(parsed.query)
            query = (params.get("q") or [""])[0]
            limit = int((params.get("limit") or ["20"])[0])
            category = (params.get("category") or [""])[0] or None
            kind = (params.get("kind") or params.get("evidenceType") or [""])[0] or None
            min_score_text = (params.get("minScore") or params.get("minConfidence") or [""])[0]
            try:
                min_score = int(float(min_score_text)) if min_score_text else None
            except ValueError:
                min_score = None
            self._json({
                "ok": True,
                **search_knowledge(query, min(max(limit, 1), 50), category, kind, min_score),
                "filters": {
                    "category": category,
                    "kind": kind,
                    "minScore": min_score
                }
            })
            return
        if route == "/api/stations/search":
            params = parse_qs(parsed.query)
            query = (params.get("q") or [""])[0]
            try:
                limit = int((params.get("limit") or ["12"])[0])
            except ValueError:
                limit = 12
            self._json({"ok": True, **search_stations(query, limit)})
            return
        if route == "/api/stations/context":
            params = parse_qs(parsed.query)
            name = (params.get("name") or params.get("q") or [""])[0]
            self._json(station_context_payload(name))
            return
        if route == "/api/research/benchmark-cases":
            self._json({"ok": True, **benchmark_cases_status()})
            return
        if route == "/api/bootstrap":
            self._json({
                "factors": FACTORS,
                "rules": RULES,
                "stations": STATIONS,
                "ridership": RIDERSHIP,
                "ridershipForecast": RIDERSHIP_FORECAST,
                "stationMemory": {
                    "version": load_station_memory_data().get("version"),
                    "records": list_station_memory_records(),
                },
                "stationOperations": STATION_OPERATIONS,
                "stationAmenities": STATION_AMENITIES,
                "inputSchema": INPUT_SCHEMA,
                "pptRules": PPT_RULES,
                "knowledgeCatalog": KNOWLEDGE_CATALOG,
                "sourceManifest": source_manifest_records(),
                "unparsedSources": UNPARSED_SOURCES,
                "platformCapabilities": build_platform_capability_status(),
                "demos": DEMOS,
                "projects": list_project_summaries(),
                "exports": list_export_files(20)
            })
            return
        if route == "/api/sources":
            self._json({
                "ok": True,
                "catalog": KNOWLEDGE_CATALOG,
                "sources": source_manifest_records(),
                "unparsedSources": UNPARSED_SOURCES
            })
            return
        if route == "/api/exports":
            self._json({"ok": True, "exports": list_export_files(80)})
            return
        if route == "/api/schematic/user-geometry":
            self._schematic_user_geometry()
            return
        if route == "/api/delivery/manifest":
            self._json({"ok": True, **build_delivery_manifest_payload()})
            return
        if route == "/api/delivery/package":
            filename, data = build_delivery_package()
            self._binary(data, "application/zip", filename)
            return
        if route == "/api/projects":
            self._json({"ok": True, "projects": list_project_summaries()})
            return
        if route.startswith("/api/projects/"):
            record_id = unquote(route.rsplit("/", 1)[-1])
            record = find_project_record(record_id)
            if not record:
                self.send_error(HTTPStatus.NOT_FOUND, "Project not found")
                return
            self._json({"ok": True, "record": record})
            return
        if route.startswith("/exports/"):
            self._export_static(route)
            return
        self._static(route)

    def do_POST(self) -> None:
        route = urlparse(self.path).path
        if route == "/api/schematic/user-geometry":
            self._save_schematic_user_geometry()
            return
        if route == "/api/schematic/export-png":
            self._export_schematic_png()
            return
        if route in {"/api/generated-images", "/api/admin/station-outlines", "/api/admin/station-outlines/apply", "/api/station-memory", "/api/station-memory/apply"}:
            try:
                payload = self._read_json_body()
                if route == "/api/generated-images":
                    result = generated_image_response(payload)
                    status = HTTPStatus.OK
                    if not result.get("ok"):
                        code = (result.get("error") or {}).get("code")
                        status = HTTPStatus.BAD_GATEWAY if code == "provider_failure" else HTTPStatus.NOT_IMPLEMENTED
                    self._json(result, status=status)
                    return
                if route == "/api/station-memory":
                    record = save_station_memory_record(payload)
                    self._json({"ok": True, "record": record, "records": list_station_memory_records(record.get("identity", {}).get("canonicalName"))})
                    return
                if route == "/api/station-memory/apply":
                    self._json(apply_station_memory_payload(payload))
                    return
                if route == "/api/admin/station-outlines":
                    record = save_admin_station_outline(payload)
                    self._json({"ok": True, "record": record, "records": list_admin_station_outlines(record["stationName"])})
                    return
                result = apply_admin_station_outline_to_geometry(
                    payload.get("geometry") or {},
                    payload.get("stationName") or "",
                    payload.get("outlineId"),
                )
                self._json({"ok": result.get("applied", False), **result})
                return
            except Exception as exc:  # noqa: BLE001
                self._json({"ok": False, "error": str(exc)}, status=HTTPStatus.BAD_REQUEST)
                return
        if route not in {"/api/evaluate", "/api/projects", "/api/export", "/api/ownership/migrate"}:
            self.send_error(HTTPStatus.NOT_FOUND, "Not found")
            return
        try:
            body = self.rfile.read(int(self.headers.get("Content-Length", "0") or 0))
            payload = json.loads(body.decode("utf-8") or "{}")
            if route == "/api/evaluate":
                project, research_options = resolve_project_payload(payload)
                result = evaluate_project(project, research_options)
                self._json({"ok": True, "result": result})
                return
            if route == "/api/projects":
                project, research_options = resolve_project_payload(payload)
                record = save_project_record(project, research_options)
                self._json({"ok": True, "record": record, "projects": list_project_summaries()})
                return
            if route == "/api/ownership/migrate":
                record = migrate_project_owner(str(payload.get("projectId") or payload.get("id") or ""), payload.get("owner"))
                self._json({"ok": True, "record": record, "projects": list_project_summaries()})
                return
            result = resolve_export_result(payload)
            export = export_report_file(result)
        except Exception as exc:  # noqa: BLE001
            self._json({"ok": False, "error": str(exc)}, status=HTTPStatus.BAD_REQUEST)
            return
        self._json({"ok": True, "export": export})

    def do_DELETE(self) -> None:
        route = urlparse(self.path).path
        if not route.startswith("/api/projects/"):
            self.send_error(HTTPStatus.NOT_FOUND, "Not found")
            return
        record_id = unquote(route.rsplit("/", 1)[-1])
        if not delete_project_record(record_id):
            self.send_error(HTTPStatus.NOT_FOUND, "Project not found")
            return
        self._json({"ok": True, "projects": list_project_summaries()})

    def log_message(self, fmt: str, *args: object) -> None:
        sys.stdout.write((fmt % args) + "\n")

    def _cors(self) -> None:
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET,POST,DELETE,OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def _json(self, payload: dict, status: HTTPStatus = HTTPStatus.OK) -> None:
        data = json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
        accepts_gzip = "gzip" in (self.headers.get("Accept-Encoding") or "").lower()
        if accepts_gzip:
            data = gzip.compress(data, compresslevel=6)
        self.send_response(status)
        self._cors()
        self.send_header("Content-Type", "application/json; charset=utf-8")
        if accepts_gzip:
            self.send_header("Content-Encoding", "gzip")
            self.send_header("Vary", "Accept-Encoding")
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-store")
        self.end_headers()
        self.wfile.write(data)

    def _binary(self, data: bytes, content_type: str, filename: str | None = None) -> None:
        self.send_response(HTTPStatus.OK)
        self._cors()
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-store")
        if filename:
            self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.end_headers()
        self.wfile.write(data)

    def _read_json_body(self) -> dict:
        body = self.rfile.read(int(self.headers.get("Content-Length", "0") or 0))
        if not body:
            return {}
        return json.loads(body.decode("utf-8") or "{}")

    def _request_base_url(self) -> str:
        proto = (self.headers.get("X-Forwarded-Proto") or "http").split(",")[0].strip() or "http"
        host = (
            self.headers.get("X-Forwarded-Host")
            or self.headers.get("Host")
            or os.environ.get("HOST")
            or "127.0.0.1:8765"
        )
        host = str(host).split(",")[0].strip()
        if ":" not in host and os.environ.get("PORT"):
            host = f"{host}:{os.environ['PORT']}"
        return f"{proto}://{host}"

    def _schematic_user_geometry(self) -> None:
        if not SCHEMATIC_GEOMETRY_PATH.exists():
            self._json({"ok": False, "geometry": None})
            return
        self._binary(SCHEMATIC_GEOMETRY_PATH.read_bytes(), "application/json; charset=utf-8")

    def _save_schematic_user_geometry(self) -> None:
        try:
            payload = self._read_json_body()
            if not isinstance(payload, dict) or "parcel" not in payload:
                raise ValueError("geometry payload must include parcel")
            payload["owner"] = payload.get("owner") or current_owner_metadata("schematic_geometry")
            SCHEMATIC_DIR.mkdir(parents=True, exist_ok=True)
            SCHEMATIC_GEOMETRY_PATH.write_text(
                json.dumps(payload, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            self._json({"ok": True, "path": str(SCHEMATIC_GEOMETRY_PATH)})
        except Exception as exc:  # noqa: BLE001
            self._json({"ok": False, "error": str(exc)}, status=HTTPStatus.BAD_REQUEST)

    def _export_schematic_png(self) -> None:
        try:
            SCHEMATIC_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
            env = os.environ.copy()
            env["AMAP_EXPORT_URL"] = env.get(
                "AMAP_EXPORT_URL",
                f"{self._request_base_url()}/schematic/index.html?view=3d&export=1",
            )
            node_modules = Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "node" / "node_modules"
            if node_modules.exists() and not env.get("NODE_PATH"):
                env["NODE_PATH"] = str(node_modules)
            completed = subprocess.run(
                ["node", str(SCHEMATIC_DIR / "export_current_view.cjs")],
                cwd=str(SCHEMATIC_DIR),
                env=env,
                check=False,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=75,
            )
            if completed.returncode != 0:
                self._json(
                    {
                        "ok": False,
                        "error": subprocess_output_text(
                            completed,
                            f"PNG export process exited with code {completed.returncode} and produced no output",
                        ),
                    },
                    status=HTTPStatus.INTERNAL_SERVER_ERROR,
                )
                return
            stdout = completed.stdout or ""
            if not stdout.strip() and (SCHEMATIC_DIR / "last_export.json").exists():
                stdout = (SCHEMATIC_DIR / "last_export.json").read_text(encoding="utf-8")
            self._json(schematic_png_export_response(json.loads(stdout)))
        except Exception as exc:  # noqa: BLE001
            self._json({"ok": False, "error": str(exc)}, status=HTTPStatus.INTERNAL_SERVER_ERROR)

    def _static(self, route: str) -> None:
        if route in {"", "/"}:
            target = FRONTEND_DIR / "index.html"
        else:
            relative = unquote(route.lstrip("/"))
            target = (FRONTEND_DIR / relative).resolve()
            if not str(target).startswith(str(FRONTEND_DIR.resolve())):
                self.send_error(HTTPStatus.FORBIDDEN, "Forbidden")
                return
        if not target.exists() or not target.is_file():
            self.send_error(HTTPStatus.NOT_FOUND, "Not found")
            return
        mime = mimetypes.guess_type(str(target))[0] or "application/octet-stream"
        if target == (SCHEMATIC_DIR / "index.html").resolve():
            key, security_code = amap_credentials()
            html = target.read_text(encoding="utf-8")
            html = html.replace("__AMAP_JS_KEY__", key)
            html = html.replace("__AMAP_SECURITY_CODE__", security_code)
            data = html.encode("utf-8")
            mime = "text/html; charset=utf-8"
        else:
            data = target.read_bytes()
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", mime)
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-store")
        self.end_headers()
        self.wfile.write(data)

    def _export_static(self, route: str) -> None:
        relative = unquote(route.removeprefix("/exports/"))
        target = (EXPORT_DIR / relative).resolve()
        if not str(target).startswith(str(EXPORT_DIR.resolve())):
            self.send_error(HTTPStatus.FORBIDDEN, "Forbidden")
            return
        if not target.exists() or not target.is_file():
            self.send_error(HTTPStatus.NOT_FOUND, "Export not found")
            return
        data = target.read_bytes()
        mime = mimetypes.guess_type(str(target))[0] or "application/octet-stream"
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", mime)
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Content-Disposition", f'attachment; filename="{target.name}"')
        self.end_headers()
        self.wfile.write(data)


def main() -> None:
    config = server_config()
    host = config["host"]
    port = config["port"]
    if "--host" in sys.argv:
        host = sys.argv[sys.argv.index("--host") + 1]
    if "--port" in sys.argv:
        port = int(sys.argv[sys.argv.index("--port") + 1])
    server = ThreadingHTTPServer((host, port), Handler)
    print(f"http://{host}:{port}/", flush=True)
    server.serve_forever()


if __name__ == "__main__":
    main()
