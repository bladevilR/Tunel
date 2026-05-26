from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "互联互通智能体运营维护及算力费用估算报告_汇报版_20260518.docx"


ACCENT = "1F4E79"
ACCENT_LIGHT = "D9EAF7"
HEADER_FILL = "EEF3F8"
CALLOUT_FILL = "F4F8FB"
BORDER = "B7C9D9"
MUTED = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_font(run, bold=False, size=11, color=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(level=level)
    p.text = ""
    run = p.add_run(text)
    set_font(run, bold=True, size={1: 16, 2: 13, 3: 12}.get(level, 11), color=ACCENT if level <= 2 else "1F3A5F")
    set_keep_with_next(p)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [16.2])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_border(cell, color="A8C4DC", size="8")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, bold=True, size=11.5, color=ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float], font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths_cm)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, bold=True, size=font_size, color=ACCENT)
    for row_values in rows:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_border(cell)
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == len(row_values) - 1 and len(text) < 15 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, size=font_size)
    doc.add_paragraph()
    return table


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, "1F3A5F", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Cm(0.7)
        style.paragraph_format.first_line_indent = Cm(-0.35)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.12

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("互联互通智能体 | 运营维护及算力费用估算")
    set_font(r, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("内部测算与立项汇报参考")
    set_font(r, size=9, color="666666")
    return doc


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("互联互通智能体")
    set_font(r, bold=True, size=24, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("运营维护及算力费用估算报告")
    set_font(r, bold=True, size=22, color=ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("立项汇报版")
    set_font(r, size=14, color="555555")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.2, 8.8])
    info = [
        ("项目名称", "苏州轨道交通站点周边互联互通智能体"),
        ("报告用途", "立项汇报、年度运营维护费用测算"),
        ("估算期间", "运营后第一年"),
        ("版本日期", "2026 年 5 月 18 日"),
    ]
    for i, (k, v) in enumerate(info):
        for j, text in enumerate((k, v)):
            cell = table.cell(i, j)
            set_cell_border(cell)
            set_cell_margins(cell, top=120, bottom=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if j == 0:
                set_cell_shading(cell, HEADER_FILL)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            set_font(r, bold=(j == 0), size=10.5, color=ACCENT if j == 0 else None)

    add_callout(
        doc,
        "汇报主口径",
        "第一年可明确测算的直接资源及接口费用建议按 1-6 万元/年预留；人工运维、数据维护、报告模板优化、培训及专项支持等服务类费用不做固定金额估算，按后续明确的服务范围和协作机制另行协定。",
    )
    doc.add_page_break()


def build_doc() -> None:
    doc = setup_doc()
    add_cover(doc)

    add_heading(doc, "一、汇报摘要", 1)
    add_callout(
        doc,
        "核心结论",
        "本项目第一年成本重点不是自建算力，而是系统稳定运行、数据持续更新、知识库维护、报告口径优化和专项支持。建议采用“轻量服务器 + 可替换大模型 API + 轻量运维服务”的方式，第一年按 10-18 万元/年作为主预算口径。",
    )
    add_table(
        doc,
        ["类别", "建议对外表述", "年度金额"],
        [
            ["基础运行资源", "用于系统部署、存储备份、日志、安全配置等基础环境", "0.5-1.5 万元"],
            ["智能模型与接口", "用于综合研判、报告生成、联网证据补强；采用 API 按量计费并可封顶", "0.5-2 万元"],
            ["基础运维服务", "用于服务巡检、异常处理、版本小更新和备份恢复演练", "1-2.5 万元"],
            ["数据与知识库维护", "用于站点、客流、出入口、周边配套、设计规则和知识库资料更新", "2-5 万元"],
            ["报告模板与业务规则优化", "用于试运行反馈下的评分解释、报告结构、汇报模板和审查口径优化", "2-4 万元"],
            ["培训及专项支持", "用于用户培训、重点汇报前支持和专项问题响应", "0.5-1.5 万元"],
            ["预备费", "用于接口价格波动、临时扩容和未预见事项", "1-2.5 万元"],
        ],
        [3.1, 10.0, 3.1],
        font_size=9.2,
    )

    add_heading(doc, "二、系统范围与费用边界", 1)
    add_para(
        doc,
        "根据现有交付系统，当前智能体已包含项目录入、站点数据补齐、互联互通必要性评分、联通方式推荐、知识库检索、模型主导综合研判、项目保存、报告导出、示意图表达和交付包部署等能力。",
    )
    add_table(
        doc,
        ["模块", "费用影响"],
        [
            ["后端 API 服务", "需要服务器、运行环境、日志、备份和故障维护。"],
            ["前端工作台", "需要浏览器兼容维护、交互优化和部署维护。"],
            ["项目库", "需要数据备份、版本管理和清理机制。"],
            ["知识库", "需要规则文件、设计指引、站点资料持续更新。"],
            ["报告导出", "会产生 Word/PDF/JSON/CSV 文件存储和模板维护。"],
            ["模型研判", "会产生大模型 API 调用费用；如私有化部署则产生 GPU 算力费用。"],
            ["地图/示意图", "可能涉及高德地图 Key、接口额度或商业授权。"],
            ["搜索/联网证据", "如启用外部搜索服务，会产生搜索接口费用。"],
        ],
        [4.0, 12.2],
        font_size=9.6,
    )
    add_para(
        doc,
        "现有数据规模包括约 240 条站点数据、约 235 条客流数据、约 249 条运营数据、约 235 条周边配套数据、约 259 条项目库数据，以及约 1508 条知识库分块。该规模决定了第一年更适合采用小型服务器和大模型 API 的轻量化运营方式。",
    )

    add_heading(doc, "三、估算基本假设", 1)
    add_heading(doc, "1. 部署假设", 2)
    add_table(
        doc,
        ["部署方式", "说明", "成本特征"],
        [
            ["内网/既有服务器部署", "使用业主或项目既有 Windows 服务器部署", "服务器新增成本低，主要为维护和备份成本"],
            ["云服务器部署", "采购 2-4 核、4-8GB 内存的小型云服务器", "有持续服务器、磁盘、备份、安全和网络费用"],
        ],
        [4.0, 6.2, 6.0],
        font_size=9.4,
    )
    add_para(doc, "当前系统为 Python 标准库 HTTP 服务、HTML/CSS/JavaScript 前端、JSON 数据文件和 Word/PDF 导出，不需要 GPU 服务器即可运行。")

    add_heading(doc, "2. 模型调用假设", 2)
    add_para(doc, "第一年建议接入主流大模型 API，并保持供应商可替换。可选模型包括通义千问、DeepSeek、Kimi、豆包等，实际以效果、稳定性、内网合规和采购便利性综合选择。")
    add_bullets(
        doc,
        [
            "模型调用主要用于项目综合研判、规则评分与模型结论差异解释、报告正文生成与润色。",
            "重点项目会产生多轮重跑、专家复核、报告模板调试和汇报前专项优化。",
            "可通过模型分级调用、缓存、摘要、额度上限和异常调用告警控制费用。",
        ],
    )
    add_table(
        doc,
        ["档位", "年项目量", "每项目调用", "单轮输入", "单轮输出", "适用场景"],
        [
            ["低用量", "300 个", "2 轮", "12 万 Token", "2.5 万 Token", "小范围试点、少量汇报"],
            ["中用量", "600 个", "3 轮", "18 万 Token", "4 万 Token", "常规试运行、项目批量评估"],
            ["高用量", "1200 个", "5 轮", "25 万 Token", "8 万 Token", "全量推广、频繁重跑、重点项目深度研判"],
        ],
        [2.3, 2.3, 2.4, 2.7, 2.7, 3.8],
        font_size=8.9,
    )

    add_heading(doc, "四、第一年费用估算", 1)
    add_heading(doc, "1. 内部管理口径", 2)
    add_para(doc, "内部管理口径用于预算测算、成本控制和采购策略设计，建议保留较细的费用拆分。")
    add_table(
        doc,
        ["成本项", "估算依据", "第一年估算"],
        [
            ["服务器/部署环境", "2-4 核、4-8GB 内存服务器；若使用既有内网服务器则仅考虑维护和备份", "0.5-1.5 万元"],
            ["存储与备份", "项目库、导出报告、日志、知识库文件、备份快照；按小规模文件型系统预留", "0.2-0.8 万元"],
            ["大模型 API 调用", "按 300-1200 个项目/年、2-5 轮/项目测算；预算含重跑、备用模型和强模型余量", "0.5-2.0 万元"],
            ["知识库向量化/检索", "新资料入库、Embedding、知识库切分、检索调优；如仅使用本地文本检索则可降低", "0.1-0.5 万元"],
            ["地图/外部接口", "高德地图 Key、搜索服务、联网证据补强；商业授权或专用服务包另计", "0-1.0 万元"],
            ["安全与账号管理", "API Key 管理、额度告警、访问控制、证书、日志审计", "0.2-1.0 万元"],
            ["基础技术运维", "按月度巡检、备份确认、异常重启、依赖小版本检查等轻量事项预留；内部人员处理时可进一步压缩，若采购供应商 SLA 应另按年度服务包列项", "0.6-2.0 万元"],
            ["业务数据维护", "按季度或半年度更新站点、客流、TOD、出入口、周边配套、知识库资料，并完成数据校核和入库", "2.0-5.0 万元"],
            ["模型与报告模板迭代", "第一年前 6-12 个月按反馈调整提示词、报告结构、评分解释和导出格式", "1.5-4.0 万元"],
            ["培训与应急支持", "用户培训、汇报前支持、重点项目问题响应", "0.5-1.5 万元"],
            ["预备费", "价格波动、临时扩容、接口变更、专项调优", "1.0-3.0 万元"],
        ],
        [3.3, 9.1, 3.8],
        font_size=8.7,
    )
    add_callout(doc, "内部建议", "第一年内部预算可按 10-18 万元/年预留；若要求更充分的供应商响应、重点项目陪跑和多轮报告优化，可内部按 20-25 万元/年设置增强上限。")

    add_heading(doc, "2. 对外汇报口径", 2)
    add_para(doc, "对外汇报口径建议聚焦费用类型、预算区间和成本可控性，不展开具体供应商单价、Token 细项和人天单价。")
    add_table(
        doc,
        ["费用类别", "对外说明", "年度建议金额"],
        [
            ["基础运行资源费", "用于系统部署、运行环境、存储备份、日志、安全配置等基础资源", "0.5-1.5 万元"],
            ["智能模型与接口调用费", "用于综合研判、报告生成、联网证据补强等智能化能力调用；采用 API 方式，按量计费、可控封顶", "0.5-2 万元"],
            ["基础运维服务费", "用于服务巡检、异常处理、版本小更新、备份恢复演练；重度驻场或 SLA 服务另计", "1-2.5 万元"],
            ["数据与知识库维护费", "用于站点、客流、出入口、周边配套、设计规则、知识库资料的持续更新", "2-5 万元"],
            ["报告模板与业务规则优化费", "用于根据试运行反馈优化评分解释、报告结构、汇报模板和审查口径", "2-4 万元"],
            ["培训及专项支持费", "用于用户培训、重点汇报前支持和专项问题响应", "0.5-1.5 万元"],
            ["预备费", "用于接口价格变动、临时扩容和未预见事项", "1-2.5 万元"],
        ],
        [3.5, 9.1, 3.6],
        font_size=8.9,
    )
    add_callout(doc, "对外建议口径", "第一年运营维护及算力相关费用建议按 10-18 万元/年测算，其中基础云资源和智能模型接口费用约 1-4 万元，系统运维、数据维护和业务规则优化约 7-12 万元。若后续采购供应商 SLA、驻场支持或全量项目陪跑，可作为增强服务另列。")

    add_heading(doc, "五、模型费用测算说明", 1)
    add_para(doc, "模型费用本质上按 Token 或调用量计费。公开价格显示，主流大模型 API 已具备较低的按量调用成本，但第一年预算应考虑反复重跑、重点项目强模型调用、备用供应商和调用审计等管理因素。")
    add_table(
        doc,
        ["模型参考", "公开单价口径", "说明"],
        [
            ["通义千问 qwen-plus", "中国内地非思考模式，0-128K Token 区间，输入约 0.8 元/百万 Token，输出约 2 元/百万 Token", "适合常规研判和报告生成"],
            ["DeepSeek chat", "输入缓存未命中约 0.27 美元/百万 Token，输出约 1.10 美元/百万 Token", "适合成本敏感场景，可作为备选"],
            ["更强/长上下文模型", "输入输出单价更高，长上下文或思考模式成本上升明显", "适合重点项目复核，不宜全量滥用"],
        ],
        [3.6, 8.6, 4.0],
        font_size=8.8,
    )
    add_para(doc, "按中用量场景估算：600 个项目/年，每项目 3 轮调用，单轮输入 18 万 Token、输出 4 万 Token，年输入量约 3.24 亿 Token，年输出量约 7200 万 Token。若采用常规经济型模型，裸 Token 成本通常在数百元至数千元级别；预算建议保留 0.5-2 万元，是为覆盖重跑、调试、强模型、备用模型和调用治理余量。")

    add_heading(doc, "六、为什么不建议第一年自建 GPU 算力", 1)
    add_para(doc, "本系统的核心任务是互联互通评估、知识库辅助研判和报告生成，不属于高并发大模型在线服务。第一年更重要的是验证业务流程、数据质量、报告口径和专家复核机制。")
    add_table(
        doc,
        ["成本类型", "影响"],
        [
            ["GPU 服务器", "一次性投入高，且需考虑显存、冗余、散热、电力和机房环境。"],
            ["模型部署与调优", "需要专门算法/工程人员维护模型推理框架。"],
            ["安全与运维", "需要持续处理显卡驱动、推理服务、队列、日志、故障和升级。"],
            ["模型效果", "本地模型效果未必优于成熟 API，且需要额外评测。"],
        ],
        [4.0, 12.2],
        font_size=9.4,
    )
    add_callout(doc, "内部判断", "除非后续出现强制内网私有化、数据完全不得出域、调用量显著放大等要求，否则第一年不建议配置 GPU 算力预算。建议保留未来可私有化部署的技术路线，但本期以 API 接入为主。")

    add_heading(doc, "七、内部口径与对外口径差异", 1)
    add_table(
        doc,
        ["项目", "内部口径", "对外口径"],
        [
            ["模型供应商", "可列通义千问、DeepSeek、Kimi、豆包等备选，并比较效果、价格、合规和稳定性", "表述为“主流大模型 API，可替换接入”"],
            ["模型单价", "可按 Token 单价测算，设置调用量和额度上限", "不披露具体供应商单价，只说明按量计费、可封顶"],
            ["人工维护", "内部可按轻量人天核算；供应商 SLA、驻场响应、专项陪跑应另列服务包，不与基础巡检混算", "表述为基础运维、数据更新、业务规则优化；重度服务另列"],
            ["云资源", "可按服务器、存储、备份、安全分项测算", "表述为基础运行资源费"],
            ["地图/搜索接口", "可列 Key、接口额度、商业授权风险", "表述为外部接口服务，按实际启用情况结算"],
            ["总预算", "可保留 10-18 万基础档和 20-25 万增强档", "建议统一报 10-18 万/年，必要时说明增强服务另列"],
        ],
        [3.1, 6.7, 6.4],
        font_size=8.8,
    )

    add_heading(doc, "八、建议写入立项材料的文字", 1)
    add_heading(doc, "1. 简版表述", 2)
    add_para(doc, "系统运营后第一年主要费用包括基础运行资源、智能模型接口调用、基础运维、数据与知识库维护、报告模板和业务规则优化等。考虑到本系统以轻量化 Web 服务和大模型 API 接入为主，不建议第一年自建 GPU 算力。年度运营维护及算力相关费用建议按 10-18 万元测算，其中直接云资源和模型接口费用约 1-4 万元，系统运维、数据维护和业务规则优化约 7-12 万元。实际费用可通过按量调用、额度封顶、模型分级调用和部署资源复用进行控制。")
    add_heading(doc, "2. 稍详细表述", 2)
    add_para(doc, "本项目第一年处于试运行和业务口径固化阶段，费用重点并非单纯算力采购，而是系统稳定运行、数据持续更新、知识库维护、模型研判质量优化和报告成果迭代。系统采用可替换的大模型 API 接入方式，避免一次性投入 GPU 服务器，模型费用按调用量结算并可设置额度上限。建议第一年按 10-18 万元/年预留运营维护及算力相关费用，覆盖基础运行资源、模型接口、数据维护、系统巡检、备份恢复、报告模板优化、用户培训和专项支持。若需供应商 SLA、驻场响应或重点项目全程陪跑，可另列增强服务。后续进入稳定运行期后，可依据实际项目量和模型调用量进一步优化年度费用。")

    add_heading(doc, "九、建议预算方案", 1)
    add_table(
        doc,
        ["方案", "适用情况", "建议金额"],
        [
            ["方案 A：保守试运行", "使用既有服务器、小范围试点、项目量较少、报告模板变动不大", "8 万元"],
            ["方案 B：推荐常规运行", "第一年正式试运行、多部门使用、需要持续优化报告和数据口径", "15 万元"],
            ["方案 C：增强保障", "全量推广、高频汇报、强支持响应、多模型比选或较多专项定制", "25 万元"],
        ],
        [4.0, 8.8, 3.4],
        font_size=9.3,
    )
    add_para(doc, "若当前仅做一期立项，建议采用方案 B 作为推荐口径，方案 A 作为压缩口径，方案 C 仅作为增强服务上限预案。")

    add_heading(doc, "十、结论", 1)
    add_para(doc, "本项目第一年运营后的费用应区分直接资源费用和运维服务费用。直接资源费用包括服务器、存储、备份、模型 API、地图/搜索接口等，预计约 1-4 万元/年；运维服务费用包括系统巡检、异常处理、数据更新、知识库维护、报告模板优化、用户培训和专项支持，预计约 7-12 万元/年。")
    add_callout(doc, "最终建议", "综合考虑第一年试运行、业务规则固化和汇报成果质量要求，建议立项按 10-18 万元/年作为运营维护及算力相关费用预算。内部管理可按 20-25 万元设置增强服务余量；对外汇报建议采用 10-18 万元/年的稳妥口径，并强调采用按量调用、额度封顶、模型可替换和资源复用机制控制成本。")

    add_heading(doc, "十一、依据来源", 1)
    sources = [
        "本项目交付文档：outputs/interconnect_agent_system/docs/handover_20260514.md，包括系统结构、功能范围、数据规模、API 和环境变量说明。",
        "本项目 LLM 接入契约：outputs/interconnect_agent_system/docs/llm_integration_contract.md，包括模型研判、联网证据、报告输出和风险控制要求。",
        "阿里云百炼模型价格：https://help.aliyun.com/zh/model-studio/model-pricing。",
        "DeepSeek API 官方价格：https://api-docs.deepseek.com/quick_start/pricing-details-usd。",
        "阿里云 ECS 计费说明：https://help.aliyun.com/zh/ecs/billing-overview。",
        "阿里云 OSS 计费说明：https://www.alibabacloud.com/help/zh/oss/billing-overview。",
    ]
    add_bullets(doc, sources)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def build_doc_v2() -> None:
    doc = setup_doc()
    add_cover(doc)

    add_heading(doc, "一、汇报摘要", 1)
    add_callout(
        doc,
        "核心结论",
        "本项目第一年不建议自建 GPU 算力，建议采用轻量服务器和可替换大模型 API 的方式运行。当前仅对服务器、存储、模型接口、地图/搜索、可选知识库向量化等可按资源规格或公开单价测算的直接费用给出区间；人工运维、数据维护、报告模板优化、培训及专项支持等服务类费用不做固定金额估算，待后续明确服务范围、响应机制、交付物和责任边界后另行协定。",
    )
    add_table(
        doc,
        ["类别", "建议对外表述", "估算口径"],
        [
            ["基础运行资源", "用于系统部署、存储备份、日志、安全配置等基础环境", "0.5-1.5 万元"],
            ["智能模型与接口", "用于综合研判、报告生成、联网证据补强；采用 API 按量计费并可封顶", "0.5-2 万元"],
            ["地图/外部搜索接口", "用于地图展示、可选联网证据补强；商业授权或专用服务包另计", "0-1 万元"],
            ["知识库向量化/检索", "如启用 Embedding 或向量检索服务，按模型或云服务实际用量计费", "0.1-0.5 万元"],
            ["人工运维服务", "具体巡检频次、响应时限、故障处理边界尚未确定", "后续协定"],
            ["数据与知识库维护", "资料更新频次、数据校核方式、责任分工尚未确定", "后续协定"],
            ["报告模板与规则优化", "需根据试运行反馈、汇报要求和专家意见确定工作范围", "后续协定"],
            ["培训及专项支持", "培训对象、次数、汇报前支持方式尚未确定", "后续协定"],
        ],
        [3.2, 9.6, 3.4],
        font_size=9.0,
    )

    add_heading(doc, "二、系统范围与费用边界", 1)
    add_para(
        doc,
        "根据现有交付系统，当前智能体已包含项目录入、站点数据补齐、互联互通必要性评分、联通方式推荐、知识库检索、模型主导综合研判、项目保存、报告导出、示意图表达和交付包部署等能力。",
    )
    add_table(
        doc,
        ["模块", "费用影响"],
        [
            ["后端 API 服务", "涉及服务器、运行环境、日志和备份资源；人工维护方式待后续确定。"],
            ["项目库与导出目录", "涉及项目数据、Word/PDF/JSON/CSV 文件存储和备份。"],
            ["知识库", "涉及资料入库、检索方式和知识更新；人工维护频次待后续确定。"],
            ["模型研判", "涉及大模型 API 调用费用；如私有化部署则另行测算 GPU 算力。"],
            ["地图/示意图", "可能涉及高德地图 Key、接口额度或商业授权。"],
            ["搜索/联网证据", "如启用外部搜索服务，会产生搜索接口费用。"],
        ],
        [4.2, 12.0],
        font_size=9.4,
    )
    add_para(
        doc,
        "现有数据规模包括约 240 条站点数据、约 235 条客流数据、约 249 条运营数据、约 235 条周边配套数据、约 259 条项目库数据，以及约 1508 条知识库分块。该规模决定了第一年更适合采用小型服务器和大模型 API 的轻量化运营方式。",
    )

    add_heading(doc, "三、估算基本假设", 1)
    add_heading(doc, "1. 部署假设", 2)
    add_table(
        doc,
        ["部署方式", "说明", "直接费用特征"],
        [
            ["内网/既有服务器部署", "使用业主或项目既有 Windows 服务器部署", "新增服务器费用低，主要为备份、存储或少量配置资源"],
            ["云服务器部署", "采购 2-4 核、4-8GB 内存的小型云服务器", "有服务器、磁盘、备份、安全和网络费用"],
        ],
        [4.0, 6.2, 6.0],
        font_size=9.3,
    )
    add_para(doc, "当前系统为 Python 标准库 HTTP 服务、HTML/CSS/JavaScript 前端、JSON 数据文件和 Word/PDF 导出，不需要 GPU 服务器即可运行。")

    add_heading(doc, "2. 模型调用假设", 2)
    add_para(doc, "第一年建议接入主流大模型 API，并保持供应商可替换。可选模型包括通义千问、DeepSeek、Kimi、豆包等，实际以效果、稳定性、内网合规和采购便利性综合选择。模型调用费用可按 Token 或调用量测算，人工服务费用不在此处折算。")
    add_table(
        doc,
        ["档位", "年项目量", "每项目调用", "单轮输入", "单轮输出", "适用场景"],
        [
            ["低用量", "300 个", "2 轮", "12 万 Token", "2.5 万 Token", "小范围试点、少量汇报"],
            ["中用量", "600 个", "3 轮", "18 万 Token", "4 万 Token", "常规试运行、项目批量评估"],
            ["高用量", "1200 个", "5 轮", "25 万 Token", "8 万 Token", "全量推广、频繁重跑、重点项目深度研判"],
        ],
        [2.3, 2.3, 2.4, 2.7, 2.7, 3.8],
        font_size=8.8,
    )

    add_heading(doc, "四、第一年费用估算", 1)
    add_heading(doc, "1. 内部管理口径", 2)
    add_para(doc, "内部管理口径仅对可按资源规格、公开接口单价或实际启用情况测算的直接费用给出金额。人工服务类费用因服务边界尚未确定，本报告不做数值估算。")
    add_table(
        doc,
        ["成本项", "估算依据", "第一年估算"],
        [
            ["服务器/部署环境", "2-4 核、4-8GB 内存服务器；若使用既有内网服务器则仅考虑新增配置或备份资源", "0.5-1.5 万元"],
            ["存储与备份", "项目库、导出报告、日志、知识库文件、备份快照；按小规模文件型系统预留", "0.2-0.8 万元"],
            ["大模型 API 调用", "按 300-1200 个项目/年、2-5 轮/项目测算；预算含重跑、备用模型和强模型余量", "0.5-2.0 万元"],
            ["知识库向量化/检索", "如启用 Embedding 或向量检索服务，按模型或云服务实际用量计费；如仅使用本地文本检索则可不单列", "0.1-0.5 万元"],
            ["地图/外部接口", "高德地图 Key、搜索服务、联网证据补强；商业授权或专用服务包另计", "0-1.0 万元"],
            ["人工运维服务", "服务频次、响应时限、故障处理边界、是否驻场或 SLA 尚未确定", "后续协定"],
            ["业务数据维护", "站点、客流、TOD、出入口、周边配套、知识库资料的更新频次和责任边界尚未确定", "后续协定"],
            ["模型与报告模板迭代", "提示词、报告结构、评分解释、导出格式等优化范围需根据试运行反馈确定", "后续协定"],
            ["培训与专项支持", "培训次数、培训对象、汇报前支持和重点项目响应方式尚未确定", "后续协定"],
        ],
        [3.3, 9.1, 3.8],
        font_size=8.7,
    )
    add_callout(doc, "内部建议", "当前可明确测算的直接资源及接口费用约 1-6 万元/年。人工运维、数据维护、报告模板优化、培训及专项支持等服务类费用不在本报告内固定估算，建议在后续明确服务内容、频次、交付物、响应机制后另行协定。")

    add_heading(doc, "2. 对外汇报口径", 2)
    add_para(doc, "对外汇报口径建议只区分“可按量测算的资源及接口费用”和“后续协定的服务类费用”，避免在运维方式尚未确定时提前锁定人工费用。")
    add_table(
        doc,
        ["费用类别", "对外说明", "年度建议金额"],
        [
            ["基础运行资源费", "用于系统部署、运行环境、存储备份、日志、安全配置等基础资源", "0.5-1.5 万元"],
            ["智能模型与接口调用费", "用于综合研判、报告生成、联网证据补强等智能化能力调用；采用 API 方式，按量计费、可控封顶", "0.5-2 万元"],
            ["地图/外部搜索接口费", "用于地图展示、可选外部搜索或联网证据补强；按实际启用情况结算", "0-1 万元"],
            ["知识库向量化/检索费", "如启用向量化或检索服务，按实际用量结算；不启用则不单列", "0.1-0.5 万元"],
            ["运营维护服务费", "服务范围、维护方式、响应机制、交付物尚未确定", "后续协定"],
            ["数据与知识库维护费", "资料更新频次、数据校核方式、责任分工尚未确定", "后续协定"],
            ["报告模板与业务规则优化费", "根据试运行反馈和汇报要求确定优化范围", "后续协定"],
            ["培训及专项支持费", "根据培训对象、次数和汇报支持需求另行确定", "后续协定"],
        ],
        [3.5, 9.1, 3.6],
        font_size=8.9,
    )
    add_callout(doc, "对外建议口径", "第一年可测算的直接资源及接口费用建议按 1-6 万元/年预留；运营维护、数据维护、报告模板优化、培训及专项支持等人工服务类费用，按照后续确定的服务范围和协作机制另行协定。")

    add_heading(doc, "五、模型费用测算说明", 1)
    add_para(doc, "模型费用本质上按 Token 或调用量计费。公开价格显示，主流大模型 API 已具备较低的按量调用成本，但第一年预算应考虑反复重跑、重点项目强模型调用、备用供应商和调用审计等管理因素。")
    add_table(
        doc,
        ["模型参考", "公开单价口径", "说明"],
        [
            ["通义千问 qwen-plus", "中国内地非思考模式，0-128K Token 区间，输入约 0.8 元/百万 Token，输出约 2 元/百万 Token", "适合常规研判和报告生成"],
            ["DeepSeek chat", "输入缓存未命中约 0.27 美元/百万 Token，输出约 1.10 美元/百万 Token", "适合成本敏感场景，可作为备选"],
            ["更强/长上下文模型", "输入输出单价更高，长上下文或思考模式成本上升明显", "适合重点项目复核，不宜全量滥用"],
        ],
        [3.6, 8.6, 4.0],
        font_size=8.8,
    )
    add_para(doc, "按中用量场景估算：600 个项目/年，每项目 3 轮调用，单轮输入 18 万 Token、输出 4 万 Token，年输入量约 3.24 亿 Token，年输出量约 7200 万 Token。若采用常规经济型模型，裸 Token 成本通常在数百元至数千元级别；预算建议保留 0.5-2 万元，是为覆盖重跑、调试、强模型、备用模型和调用治理余量。")

    add_heading(doc, "六、为什么不建议第一年自建 GPU 算力", 1)
    add_para(doc, "本系统的核心任务是互联互通评估、知识库辅助研判和报告生成，不属于高并发大模型在线服务。第一年更重要的是验证业务流程、数据质量、报告口径和专家复核机制。")
    add_table(
        doc,
        ["成本类型", "影响"],
        [
            ["GPU 服务器", "一次性投入高，且需考虑显存、冗余、散热、电力和机房环境。"],
            ["模型部署与调优", "需要专门算法/工程人员维护模型推理框架。"],
            ["安全与运维", "需要持续处理显卡驱动、推理服务、队列、日志、故障和升级。"],
            ["模型效果", "本地模型效果未必优于成熟 API，且需要额外评测。"],
        ],
        [4.0, 12.2],
        font_size=9.4,
    )
    add_callout(doc, "内部判断", "除非后续出现强制内网私有化、数据完全不得出域、调用量显著放大等要求，否则第一年不建议配置 GPU 算力预算。建议保留未来可私有化部署的技术路线，但本期以 API 接入为主。")

    add_heading(doc, "七、内部口径与对外口径差异", 1)
    add_table(
        doc,
        ["项目", "内部口径", "对外口径"],
        [
            ["模型供应商", "可列通义千问、DeepSeek、Kimi、豆包等备选，并比较效果、价格、合规和稳定性", "表述为“主流大模型 API，可替换接入”"],
            ["模型单价", "可按 Token 单价测算，设置调用量和额度上限", "不披露具体供应商单价，只说明按量计费、可封顶"],
            ["人工维护", "不在当前报告中测算固定金额；待服务范围、响应机制、交付物和责任边界明确后协定", "表述为服务类费用后续协定"],
            ["云资源", "可按服务器、存储、备份、安全分项测算", "表述为基础运行资源费"],
            ["地图/搜索接口", "可列 Key、接口额度、商业授权风险", "表述为外部接口服务，按实际启用情况结算"],
            ["总费用", "只汇总直接资源及接口费用；服务类费用不提前锁定金额", "建议表述为直接费用约 1-6 万元/年，服务类费用后续协定"],
        ],
        [3.1, 6.7, 6.4],
        font_size=8.8,
    )

    add_heading(doc, "八、建议写入立项材料的文字", 1)
    add_heading(doc, "1. 简版表述", 2)
    add_para(doc, "系统运营后第一年可明确测算的费用主要包括基础运行资源、智能模型接口调用、地图/搜索接口和可选知识库向量化服务等，建议直接资源及接口费用按 1-6 万元/年预留。人工运维、数据维护、报告模板优化、培训及专项支持等服务类费用，因维护方式、服务频次和责任边界尚未确定，建议按照后续协定另行确定。")
    add_heading(doc, "2. 稍详细表述", 2)
    add_para(doc, "本项目第一年处于试运行和业务口径固化阶段，不建议投入自建 GPU 算力，宜采用可替换的大模型 API 接入方式，模型费用按调用量结算并设置额度上限。根据当前系统规模，服务器、存储、模型接口、地图/搜索和可选知识库向量化等直接费用可按 1-6 万元/年预留。涉及人工运维、数据与知识库维护、报告模板和业务规则优化、培训及专项支持的费用，需在后续明确服务范围、响应机制、交付物和责任边界后另行协定。")

    add_heading(doc, "九、直接资源费用预留方案", 1)
    add_table(
        doc,
        ["方案", "适用情况", "直接资源及接口费用"],
        [
            ["方案 A：保守试运行", "使用既有服务器、小范围试点、项目量较少、外部接口启用较少", "1-2 万元"],
            ["方案 B：推荐常规运行", "正式试运行、多部门使用、模型调用和报告导出较稳定", "2-4 万元"],
            ["方案 C：增强预留", "项目量较高、启用外部搜索或更强模型、接口调用余量要求较高", "4-6 万元"],
        ],
        [4.0, 8.8, 3.4],
        font_size=9.3,
    )
    add_para(doc, "以上仅为直接资源及接口费用预留，不包含人工运维、数据维护、报告模板优化、培训和专项支持等服务类费用。服务类费用建议在后续明确维护方式和服务边界后另行协定。")

    add_heading(doc, "十、结论", 1)
    add_para(doc, "本项目第一年运营后的费用应区分直接资源及接口费用和人工服务类费用。直接资源及接口费用包括服务器、存储、备份、模型 API、地图/搜索接口和可选知识库向量化服务，可按 1-6 万元/年预留。人工运维、数据维护、报告模板优化、培训和专项支持等服务类费用，因具体维护方式和服务边界尚未确定，本报告不做固定金额估算。")
    add_callout(doc, "最终建议", "立项汇报中建议采用“直接资源及接口费用约 1-6 万元/年，人工服务类费用按后续协定另行确定”的口径；同时说明系统采用按量调用、额度封顶、模型可替换和资源复用机制控制成本，不建议第一年自建 GPU 算力。")

    add_heading(doc, "十一、依据来源", 1)
    add_bullets(
        doc,
        [
            "本项目交付文档：outputs/interconnect_agent_system/docs/handover_20260514.md，包括系统结构、功能范围、数据规模、API 和环境变量说明。",
            "本项目 LLM 接入契约：outputs/interconnect_agent_system/docs/llm_integration_contract.md，包括模型研判、联网证据、报告输出和风险控制要求。",
            "阿里云百炼模型价格：https://help.aliyun.com/zh/model-studio/model-pricing。",
            "DeepSeek API 官方价格：https://api-docs.deepseek.com/quick_start/pricing-details-usd。",
            "阿里云 ECS 计费说明：https://help.aliyun.com/zh/ecs/billing-overview。",
            "阿里云 OSS 计费说明：https://www.alibabacloud.com/help/zh/oss/billing-overview。",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def build_doc_v3() -> None:
    doc = setup_doc()
    add_cover(doc)

    add_heading(doc, "一、结论口径", 1)
    add_callout(
        doc,
        "建议写法",
        "本项目第一年不建议自建 GPU 算力，建议采用“集团内网资源优先、外部云资源备选、模型 API 按量调用”的方式。可明确测算的直接资源及接口费用建议按 1-6 万元/年预留；人工运维、数据维护、报告模板优化、培训及专项支持等服务类费用，按后续协定另行确定。",
    )

    add_heading(doc, "二、部署资源费用", 1)
    add_table(
        doc,
        ["部署方式", "资源安排", "费用口径"],
        [
            [
                "仅内网使用",
                "优先申请集团或业主既有服务器、存储、备份和网络访问资源；系统部署在内网环境，原则上不新增外部云服务器租赁费用。",
                "不单列服务器租赁费；如需新增内网主机、存储扩容或备份策略，按集团内部资源申请和信息化管理要求办理。",
            ],
            [
                "需要外网访问",
                "租用外部云服务器，通常包括 2-4 核 CPU、4-8GB 内存、50-100GB 云盘、备份/快照及必要公网带宽或访问流量。",
                "约 0.5-1.5 万元/年；若需更高安全等级、专线、堡垒机、等保、安全加固或公网流量较大，另行测算。",
            ],
        ],
        [3.0, 7.4, 5.8],
        font_size=9.0,
    )

    add_heading(doc, "三、模型及接口费用", 1)
    add_table(
        doc,
        ["费用项", "说明", "费用口径"],
        [
            ["大模型 API 调用", "用于项目综合研判、报告生成、重点项目重跑；采用通义千问、DeepSeek、Kimi、豆包等可替换模型。", "约 0.5-2 万元/年"],
            ["地图/外部搜索接口", "用于地图展示、可选联网证据补强；高德等商业授权或专用服务包不包含在内。", "约 0-1 万元/年"],
            ["知识库向量化/检索", "如启用 Embedding 或向量检索服务，按实际模型和云服务用量计费；仅本地文本检索时可不单列。", "约 0.1-0.5 万元/年"],
        ],
        [3.4, 9.0, 3.8],
        font_size=9.0,
    )
    add_para(doc, "上述接口费用可通过按量调用、额度封顶、模型分级调用、缓存和备用模型机制控制。模型调用本身不是第一年主要成本来源。")

    add_heading(doc, "四、后续协定事项", 1)
    add_table(
        doc,
        ["事项", "建议口径"],
        [
            ["人工运维服务", "巡检频次、响应时限、故障处理边界、是否驻场或 SLA 尚未确定，金额按后续协定。"],
            ["数据与知识库维护", "站点、客流、TOD、出入口、周边配套、知识库资料的更新频次和责任边界尚未确定，金额按后续协定。"],
            ["报告模板与规则优化", "提示词、报告结构、评分解释、导出格式等优化范围需根据试运行反馈确定，金额按后续协定。"],
            ["培训及专项支持", "培训次数、培训对象、汇报前支持和重点项目响应方式尚未确定，金额按后续协定。"],
        ],
        [4.0, 12.2],
        font_size=9.2,
    )

    add_heading(doc, "五、立项材料建议表述", 1)
    add_para(
        doc,
        "系统运营后第一年可明确测算的费用主要包括部署资源、模型接口、地图/搜索接口和可选知识库向量化服务。若系统仅在集团内网使用，服务器、存储和备份资源建议优先通过集团内部资源申请解决，不单列外部云服务器租赁费用；若需要外网访问，则需租用外部云服务器，按小型应用服务器配置估算约 0.5-1.5 万元/年。模型及接口费用建议按 1-3.5 万元/年预留。综合来看，直接资源及接口费用建议按 1-6 万元/年预留。人工运维、数据维护、报告模板优化、培训及专项支持等服务类费用，按照后续明确的服务范围和协作机制另行协定。",
    )

    add_heading(doc, "六、依据说明", 1)
    add_bullets(
        doc,
        [
            "系统规模依据：现有交付包为轻量 Web 服务，包含项目库、知识库、报告导出和模型 API 接入，不需要自建 GPU。",
            "部署费用依据：云服务器通常由实例、云盘、快照/备份、公网带宽或访问流量等计费项组成；具体金额以后续选型和采购价格为准。",
            "模型费用依据：主流大模型 API 按 Token 或调用量计费，可通过额度封顶和模型分级调用控制。",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc_v3()
