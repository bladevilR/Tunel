from __future__ import annotations

import json
import re
import zipfile
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from statistics import mean
from xml.etree import ElementTree as ET

from docx import Document
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = ROOT.parents[1]
FEEDBACK_DIR = ROOT.parent / "feedback_20260507"
DATA_DIR = ROOT / "data"
KNOWLEDGE_DIR = DATA_DIR / "knowledge"
DOCS_DIR = ROOT / "docs"

FEEDBACK_WORKBOOK = WORKSPACE_ROOT / "互联互通智能体资料缺项需求提报表-反馈版.xlsx"
AMENITIES_WORKBOOK = WORKSPACE_ROOT / "运营公司车站出入口开放及周边配套统计表202604.xlsx"

YES = "是"
NO = "否"
EMPTY_VALUES = {"", "无", "——", "-", "/", "None", "null"}


def clean(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()


def json_write(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def json_load(path: Path, default: object) -> object:
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def rel(path: Path) -> str:
    try:
        return str(path.relative_to(WORKSPACE_ROOT))
    except ValueError:
        return str(path)


def source_id(path: Path) -> str:
    text = rel(path).replace("\\", "/")
    return re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "-", text).strip("-")[:120]


def classify_source(path: Path) -> dict:
    name = path.name
    text = rel(path)
    if name.startswith("~$"):
        return {"usefulness": "无用", "priority": 0, "category": "office_temp", "reason": "Office 临时锁文件，不含业务数据"}
    if "0318" in name and path.suffix.lower() == ".pptx":
        return {"usefulness": "核心", "priority": 100, "category": "work_plan", "reason": "PPT 是用户确认的一期最高优先级规则源"}
    if "反馈版" in name:
        return {"usefulness": "核心", "priority": 98, "category": "feedback", "reason": "20260507 反馈明确评分、客流、等级和一期边界"}
    if "评价因子赋值" in name:
        return {"usefulness": "核心", "priority": 96, "category": "scoring", "reason": "评分赋值表是反馈确认的计算口径"}
    if "标准输入要素" in name:
        return {"usefulness": "核心", "priority": 94, "category": "input_schema", "reason": "标准输入界面字段来源"}
    if "日均进站" in name:
        return {"usefulness": "核心", "priority": 92, "category": "ridership", "reason": "一期客流口径来源"}
    if "TOD" in name:
        return {"usefulness": "核心", "priority": 90, "category": "station_tod", "reason": "站点 TOD 与线路预设来源"}
    if "联通接口" in name or "联通界面" in name:
        return {"usefulness": "核心", "priority": 88, "category": "station_interface", "reason": "站点出入口、联通接口、建设状态和问题来源"}
    if "运营公司车站出入口开放" in name:
        return {"usefulness": "重要", "priority": 82, "category": "station_amenities", "reason": "运营视角出入口开放、运营管理和周边配套来源"}
    if "设计导则" in text or "设计细则" in text or "设计指引" in text or "专项规划" in text or "评价体系" in text:
        return {"usefulness": "重要", "priority": 75, "category": "design_guidance", "reason": "知识库规则条文与设计参数来源，需逐步细化成规则卡片"}
    if "早期文件" in text or "豆包" in text or "仅供参考" in text:
        return {"usefulness": "低优先参考", "priority": 35, "category": "early_reference", "reason": "可参考报告表达，不能覆盖 PPT/Excel/反馈口径"}
    if path.suffix.lower() == ".png":
        return {"usefulness": "参考", "priority": 30, "category": "visual_reference", "reason": "联通意向效果图风格参考，对评分无直接作用"}
    return {"usefulness": "参考", "priority": 50, "category": "general", "reason": "资料包内补充资料"}


def chunk_text(text: str, size: int = 1200, overlap: int = 120) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        chunks.append(text[start:end].strip())
        if end == len(text):
            break
        start = max(0, end - overlap)
    return chunks


class KnowledgeBuilder:
    def __init__(self) -> None:
        self.chunks: list[dict] = []
        self.manifest: list[dict] = []
        self.unparsed: list[dict] = []

    def add_chunk(self, source: Path, title: str, text: str, kind: str, metadata: dict | None = None) -> None:
        for index, chunk in enumerate(chunk_text(text), 1):
            if not chunk:
                continue
            self.chunks.append({
                "id": f"{source_id(source)}::{kind}::{len(self.chunks) + 1}",
                "sourceId": source_id(source),
                "sourcePath": rel(source),
                "sourceName": source.name,
                "title": title,
                "kind": kind,
                "text": chunk,
                "metadata": metadata or {},
                "priority": classify_source(source)["priority"],
                "category": classify_source(source)["category"]
            })

    def add_manifest(self, source: Path, parse_status: str, counts: dict | None = None, note: str = "") -> None:
        classification = classify_source(source)
        self.manifest.append({
            "id": source_id(source),
            "path": rel(source),
            "name": source.name,
            "suffix": source.suffix.lower(),
            "size": source.stat().st_size if source.exists() else None,
            "parseStatus": parse_status,
            "usefulness": classification["usefulness"],
            "priority": classification["priority"],
            "category": classification["category"],
            "reason": classification["reason"],
            "counts": counts or {},
            "note": note
        })
        if parse_status != "parsed":
            self.unparsed.append({
                "path": rel(source),
                "parseStatus": parse_status,
                "reason": note or classification["reason"],
                "usefulness": classification["usefulness"]
            })


def read_docx(builder: KnowledgeBuilder, source: Path) -> None:
    doc = Document(source)
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for index, text in enumerate(paragraphs, 1):
        builder.add_chunk(source, f"{source.name} / 段落 {index}", text, "docx_paragraph", {"paragraph": index})
    for table_index, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", "；") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            builder.add_chunk(source, f"{source.name} / 表格 {table_index}", "\n".join(rows), "docx_table", {"table": table_index})
    builder.add_manifest(source, "parsed", {"paragraphs": len(paragraphs), "tables": len(doc.tables)})


def read_pptx(builder: KnowledgeBuilder, source: Path) -> None:
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    slide_count = 0
    with zipfile.ZipFile(source) as zf:
        names = sorted(
            [name for name in zf.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", name)],
            key=lambda name: int(re.search(r"(\d+)", name).group(1))
        )
        for slide_index, name in enumerate(names, 1):
            tree = ET.fromstring(zf.read(name))
            texts = [node.text.strip() for node in tree.findall(".//a:t", ns) if node.text and node.text.strip()]
            if texts:
                builder.add_chunk(source, f"{source.name} / 第{slide_index}页", "\n".join(texts), "ppt_slide", {"slide": slide_index})
            slide_count += 1
    builder.add_manifest(source, "parsed", {"slides": slide_count})


def read_pdf(builder: KnowledgeBuilder, source: Path) -> None:
    reader = PdfReader(str(source))
    extracted_pages = 0
    total_chars = 0
    for page_index, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            extracted_pages += 1
            total_chars += len(text.strip())
            builder.add_chunk(source, f"{source.name} / 第{page_index}页", text, "pdf_page", {"page": page_index})
    status = "parsed" if extracted_pages else "metadata_only"
    note = "" if extracted_pages else "PDF 未抽取到可用文本，可能是扫描件；需 OCR 后才能全文入库"
    builder.add_manifest(source, status, {"pages": len(reader.pages), "extractedPages": extracted_pages, "textChars": total_chars}, note)


def stringify_row(row: tuple[object, ...]) -> str:
    return " | ".join(clean(cell).replace("\n", "；") for cell in row if clean(cell))


def read_xlsx(builder: KnowledgeBuilder, source: Path) -> None:
    wb = load_workbook(source, read_only=True, data_only=True)
    sheet_summaries = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            text = stringify_row(row)
            if text:
                rows.append(text)
        sheet_summaries.append({"name": ws.title, "rows": ws.max_row, "cols": ws.max_column, "nonEmptyRows": len(rows)})
        for start in range(0, len(rows), 80):
            batch = rows[start:start + 80]
            if batch:
                builder.add_chunk(
                    source,
                    f"{source.name} / {ws.title} / 行{start + 1}-{start + len(batch)}",
                    "\n".join(batch),
                    "xlsx_rows",
                    {"sheet": ws.title, "rowStart": start + 1, "rowEnd": start + len(batch)}
                )
    wb.close()
    builder.add_manifest(source, "parsed", {"sheets": sheet_summaries})


def read_png(builder: KnowledgeBuilder, source: Path) -> None:
    with Image.open(source) as image:
        metadata = {"width": image.width, "height": image.height, "mode": image.mode}
    builder.add_chunk(source, source.name, f"联通意向效果图参考。图像尺寸：{metadata['width']}x{metadata['height']}。", "image_metadata", metadata)
    builder.add_manifest(source, "metadata_only", metadata, "图片先登记元数据和用途；若需图像语义，需要单独视觉标注")


def discover_sources() -> list[Path]:
    sources = []
    if FEEDBACK_WORKBOOK.exists():
        sources.append(FEEDBACK_WORKBOOK)
    if AMENITIES_WORKBOOK.exists():
        sources.append(AMENITIES_WORKBOOK)
    if FEEDBACK_DIR.exists():
        sources.extend(path for path in FEEDBACK_DIR.rglob("*") if path.is_file())
    return sorted(dict.fromkeys(sources), key=lambda path: (classify_source(path)["priority"] * -1, rel(path)))


def parse_sources() -> tuple[list[dict], list[dict], list[dict]]:
    builder = KnowledgeBuilder()
    for source in discover_sources():
        suffix = source.suffix.lower()
        if source.name.startswith("~$"):
            builder.add_manifest(source, "ignored", note="Office 临时锁文件，跳过")
            continue
        try:
            if suffix == ".docx":
                read_docx(builder, source)
            elif suffix == ".pptx":
                read_pptx(builder, source)
            elif suffix == ".pdf":
                read_pdf(builder, source)
            elif suffix == ".xlsx":
                read_xlsx(builder, source)
            elif suffix == ".png":
                read_png(builder, source)
            elif suffix == ".doc":
                converted = source.with_suffix(".docx")
                if converted.exists():
                    builder.add_manifest(
                        source,
                        "parsed",
                        {"convertedTo": rel(converted)},
                        note="旧版 Word 文件已转换为同名 docx，全文由转换文件入库",
                    )
                else:
                    builder.add_manifest(source, "unparsed", note="旧版二进制 Word 文件；当前环境无可靠转换器，需后续转 docx 或 OCR")
            else:
                builder.add_manifest(source, "unparsed", note=f"暂未支持的文件类型：{suffix}")
        except Exception as exc:  # noqa: BLE001
            builder.add_manifest(source, "error", note=str(exc)[:400])
    return builder.manifest, builder.chunks, builder.unparsed


def station_aliases(name: str | None) -> list[str]:
    text = clean(name)
    if not text:
        return []
    aliases = [text]
    aliases.extend(match.strip() for match in re.findall(r"[（(]([^）)]+)[）)]", text) if match.strip())
    no_paren = re.sub(r"[（(].*?[）)]", "", text).strip()
    if no_paren:
        aliases.append(no_paren)
    if no_paren.endswith("站"):
        aliases.append(no_paren[:-1])
    return list(dict.fromkeys(item for item in aliases if item))


def first_alias(name: str) -> str:
    aliases = station_aliases(name)
    if len(aliases) > 1:
        return aliases[1]
    return aliases[0] if aliases else clean(name)


def build_alias_index(records: list[dict], fields: list[str]) -> dict[str, dict]:
    index = {}
    for record in records:
        names = []
        for field in fields:
            value = record.get(field)
            if isinstance(value, list):
                names.extend(value)
            elif value:
                names.append(value)
        for name in names:
            for alias in station_aliases(name):
                index.setdefault(alias, record)
    return index


def parse_station_amenities() -> dict:
    if not AMENITIES_WORKBOOK.exists():
        return {"version": "202604", "source": "", "count": 0, "records": [], "lineSummary": []}

    wb = load_workbook(AMENITIES_WORKBOOK, read_only=True, data_only=True)
    records: dict[str, dict] = {}
    line_summary = []

    for ws in wb.worksheets:
        if ws.title == "出入口汇总":
            continue
        current_line = ""
        current_station = ""
        counts = Counter()
        stations = set()
        for row in ws.iter_rows(min_row=5, values_only=True):
            if not any(clean(cell) for cell in row):
                continue
            if len(row) > 1 and clean(row[1]):
                current_line = clean(row[1])
            if len(row) > 2 and clean(row[2]):
                current_station = clean(row[2])
            exit_no = clean(row[3] if len(row) > 3 else "")
            if not current_station or not exit_no:
                continue

            open_flag = clean(row[4] if len(row) > 4 else "")
            owner = clean(row[5] if len(row) > 5 else "")
            managed_flag = clean(row[6] if len(row) > 6 else "")
            key = first_alias(current_station)
            record = records.setdefault(key, {
                "name": key,
                "displayNames": set(),
                "aliases": set(),
                "lines": set(),
                "exitRows": 0,
                "openExitCount": 0,
                "closedExitCount": 0,
                "managedExitCount": 0,
                "owners": set(),
                "nearby": {"schools": set(), "hubs": set(), "residential": set(), "commercial": set(), "other": set()},
                "mobility": {"publicBike": [], "sharedBike": [], "parkAndRide": [], "bikeAndRide": [], "busTransfer": []},
                "sampleExits": []
            })
            record["displayNames"].add(current_station)
            for alias in station_aliases(current_station):
                record["aliases"].add(alias)
            if current_line:
                record["lines"].add(current_line)
            if owner:
                record["owners"].add(owner)
            record["exitRows"] += 1
            counts["exitRows"] += 1
            stations.add(current_station)
            if open_flag == YES:
                record["openExitCount"] += 1
                counts["open"] += 1
            elif open_flag == NO:
                record["closedExitCount"] += 1
                counts["closed"] += 1
            if managed_flag == YES:
                record["managedExitCount"] += 1
                counts["managed"] += 1

            for key_name, col in [("schools", 7), ("hubs", 8), ("residential", 9), ("commercial", 10), ("other", 11)]:
                value = clean(row[col] if len(row) > col else "")
                if value and value not in EMPTY_VALUES:
                    record["nearby"][key_name].add(value)

            if len(record["sampleExits"]) < 20:
                record["sampleExits"].append({
                    "line": current_line,
                    "exit": exit_no,
                    "isOpen": open_flag,
                    "owner": owner,
                    "isManaged": managed_flag
                })

            # These columns vary between lines, so keep compact evidence strings rather than forcing a brittle schema.
            tail_text = [clean(cell) for cell in row[12:] if clean(cell) and clean(cell) not in EMPTY_VALUES]
            joined_tail = "；".join(tail_text)
            if "公交" in joined_tail:
                record["mobility"]["busTransfer"].append(joined_tail[:160])
            if "P+R" in joined_tail or "停车" in joined_tail or "收费" in joined_tail:
                record["mobility"]["parkAndRide"].append(joined_tail[:160])
            if "自行车" in joined_tail or "单车" in joined_tail:
                record["mobility"]["publicBike"].append(joined_tail[:160])

        line_summary.append({
            "line": ws.title,
            "stationCount": len(stations),
            "exitRows": counts["exitRows"],
            "openExitCount": counts["open"],
            "closedExitCount": counts["closed"],
            "managedExitCount": counts["managed"]
        })
    wb.close()

    output_records = []
    for record in records.values():
        output_records.append({
            "name": record["name"],
            "displayNames": sorted(record["displayNames"]),
            "aliases": sorted(record["aliases"]),
            "lines": sorted(record["lines"]),
            "exitRows": record["exitRows"],
            "openExitCount": record["openExitCount"],
            "closedExitCount": record["closedExitCount"],
            "managedExitCount": record["managedExitCount"],
            "owners": sorted(record["owners"]),
            "nearby": {key: sorted(values)[:20] for key, values in record["nearby"].items()},
            "mobility": {key: values[:12] for key, values in record["mobility"].items()},
            "sampleExits": record["sampleExits"]
        })
    output_records.sort(key=lambda item: item["name"])
    payload = {
        "version": "202604",
        "source": rel(AMENITIES_WORKBOOK),
        "count": len(output_records),
        "lineSummary": line_summary,
        "records": output_records
    }
    json_write(DATA_DIR / "station_amenities.json", payload)
    return payload


def build_rule_cards() -> dict:
    factors = json_load(DATA_DIR / "factors.json", {})
    design_rules = json_load(DATA_DIR / "design_rules.json", {})
    ppt_rules = json_load(DATA_DIR / "ppt_rules_summary.json", {})

    cards = []
    for grade in factors.get("grading", []):
        cards.append({
            "id": f"grading-{grade.get('level')}",
            "type": "grading",
            "title": grade.get("level"),
            "rule": grade,
            "source": "20260507反馈 + factors.json",
            "priority": 98
        })
    for dimension in factors.get("dimensions", []):
        for factor in dimension.get("factors", []):
            cards.append({
                "id": f"factor-{factor.get('id')}",
                "type": "score_factor",
                "title": f"{dimension.get('name')} / {factor.get('name')}",
                "rule": {
                    "dimension": dimension.get("name"),
                    "factor": factor.get("name"),
                    "weight": factor.get("weight"),
                    "options": factor.get("options", []),
                    "description": factor.get("description")
                },
                "source": "评价因子赋值明细表.xlsx + factors.json",
                "priority": 96
            })
    for item in design_rules.get("connectionTypes", []):
        cards.append({
            "id": f"connection-{item.get('id')}",
            "type": "connection_type",
            "title": item.get("name"),
            "rule": item,
            "source": "设计指引类资料 + design_rules.json",
            "priority": 85
        })
    for item in design_rules.get("recommendationRules", []):
        cards.append({
            "id": f"recommendation-{item.get('id')}",
            "type": "recommendation_rule",
            "title": item.get("when"),
            "rule": item,
            "source": "PPT方案库 + design_rules.json",
            "priority": 86
        })
    for index, section in enumerate(ppt_rules.get("outputTemplateSections", []), 1):
        cards.append({
            "id": f"report-section-{index}",
            "type": "report_section",
            "title": section,
            "rule": {"order": index, "name": section},
            "source": "0318 PPT 第8页",
            "priority": 90
        })
    payload = {"version": "knowledge-20260507", "count": len(cards), "cards": cards}
    json_write(KNOWLEDGE_DIR / "rule_cards.json", payload)
    return payload


def build_station_index(amenities: dict) -> dict:
    stations = json_load(DATA_DIR / "stations.json", {"stations": []}).get("stations", [])
    ridership = json_load(DATA_DIR / "ridership.json", {"records": []}).get("records", [])
    operations = json_load(DATA_DIR / "station_operations.json", {"records": []}).get("records", [])
    amenities_records = amenities.get("records", [])

    merged: dict[str, dict] = {}

    def ensure(name: str) -> dict:
        key = first_alias(name)
        return merged.setdefault(key, {"name": key, "aliases": set(), "sources": set()})

    for station in stations:
        record = ensure(station.get("name", ""))
        record["aliases"].update(station_aliases(station.get("name", "")))
        record["sources"].add("stations.json")
        record["tod"] = station
    for item in ridership:
        record = ensure(item.get("stationName", ""))
        record["aliases"].update(station_aliases(item.get("stationName", "")))
        record["sources"].add("ridership.json")
        record["ridership"] = item
    for item in operations:
        record = ensure(item.get("name", ""))
        record["aliases"].update(item.get("aliases", []))
        record["sources"].add("station_operations.json")
        record["operations"] = item
    for item in amenities_records:
        record = ensure(item.get("name", ""))
        record["aliases"].update(item.get("aliases", []))
        record["sources"].add("station_amenities.json")
        record["amenities"] = item

    records = []
    for item in merged.values():
        item["aliases"] = sorted(item["aliases"])
        item["sources"] = sorted(item["sources"])
        records.append(item)
    records.sort(key=lambda item: item["name"])
    payload = {
        "version": "knowledge-20260507",
        "count": len(records),
        "coverage": {
            "todStations": len(stations),
            "ridershipStations": len(ridership),
            "interfaceStations": len(operations),
            "amenityStations": len(amenities_records)
        },
        "records": records
    }
    json_write(KNOWLEDGE_DIR / "station_index.json", payload)
    return payload


def build_catalog(manifest: list[dict], chunks: list[dict], unparsed: list[dict], rule_cards: dict, station_index: dict, amenities: dict) -> dict:
    source_status = Counter(item["parseStatus"] for item in manifest)
    usefulness = Counter(item["usefulness"] for item in manifest)
    categories = Counter(item["category"] for item in manifest)
    chunk_categories = Counter(item["category"] for item in chunks)
    payload = {
        "version": "knowledge-20260507",
        "builtAt": datetime.now().replace(microsecond=0).isoformat(),
        "root": rel(ROOT),
        "summary": {
            "sourceCount": len(manifest),
            "chunkCount": len(chunks),
            "ruleCardCount": rule_cards.get("count", 0),
            "stationIndexCount": station_index.get("count", 0),
            "stationAmenitiesCount": amenities.get("count", 0),
            "unparsedCount": len(unparsed)
        },
        "sourceStatus": dict(source_status),
        "sourceUsefulness": dict(usefulness),
        "sourceCategories": dict(categories),
        "chunkCategories": dict(chunk_categories),
        "files": {
            "sourceManifest": "data/knowledge/source_manifest.json",
            "knowledgeChunks": "data/knowledge/knowledge_chunks.jsonl",
            "ruleCards": "data/knowledge/rule_cards.json",
            "stationIndex": "data/knowledge/station_index.json",
            "stationAmenities": "data/station_amenities.json",
            "unparsedSources": "data/knowledge/unparsed_sources.json"
        },
        "knownLimitations": [
            "旧版 .doc 文件未全文解析，需转换为 .docx 或 PDF/OCR 后补入。",
            "部分 PDF 文本抽取量较少，可能包含扫描页；后续可做 OCR 强化。",
            "PNG 当前只入库为视觉参考元数据，未做图像语义标注。",
            "早期豆包报告按低优先参考入库，不参与规则裁决。"
        ]
    }
    json_write(KNOWLEDGE_DIR / "knowledge_catalog.json", payload)
    return payload


def write_chunks(chunks: list[dict]) -> None:
    KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
    target = KNOWLEDGE_DIR / "knowledge_chunks.jsonl"
    target.write_text("\n".join(json.dumps(chunk, ensure_ascii=False) for chunk in chunks) + "\n", encoding="utf-8")


def write_docs(catalog: dict, manifest: list[dict], unparsed: list[dict]) -> None:
    useful_lines = []
    for usefulness in ["核心", "重要", "参考", "低优先参考", "无用"]:
        grouped = [item for item in manifest if item["usefulness"] == usefulness]
        if grouped:
            useful_lines.append(f"### {usefulness}")
            for item in grouped:
                useful_lines.append(f"- `{item['path']}`：{item['reason']}（解析状态：{item['parseStatus']}）")
            useful_lines.append("")

    limitations = "\n".join(f"- {item}" for item in catalog["knownLimitations"])
    unparsed_lines = "\n".join(f"- `{item['path']}`：{item['reason']}" for item in unparsed) or "- 无"
    text = f"""# 互联互通知识数据库说明

生成时间：{catalog['builtAt']}

## 数据库构成

- 来源清单：`data/knowledge/source_manifest.json`
- 检索知识块：`data/knowledge/knowledge_chunks.jsonl`
- 规则卡片：`data/knowledge/rule_cards.json`
- 站点综合索引：`data/knowledge/station_index.json`
- 运营配套数据：`data/station_amenities.json`
- 未解析/低可用来源：`data/knowledge/unparsed_sources.json`

## 规模

- 来源文件：{catalog['summary']['sourceCount']} 个
- 知识块：{catalog['summary']['chunkCount']} 条
- 规则卡片：{catalog['summary']['ruleCardCount']} 条
- 站点索引：{catalog['summary']['stationIndexCount']} 个站点/别名主记录
- 运营配套站点：{catalog['summary']['stationAmenitiesCount']} 个

## 资料分层

{chr(10).join(useful_lines)}
## 未解析或仅元数据入库

{unparsed_lines}

## 已知限制

{limitations}
"""
    (DOCS_DIR / "knowledge_database.md").write_text(text, encoding="utf-8")


def main() -> None:
    KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
    manifest, chunks, unparsed = parse_sources()
    amenities = parse_station_amenities()
    rule_cards = build_rule_cards()
    station_index = build_station_index(amenities)
    write_chunks(chunks)
    json_write(KNOWLEDGE_DIR / "source_manifest.json", manifest)
    json_write(KNOWLEDGE_DIR / "unparsed_sources.json", unparsed)
    catalog = build_catalog(manifest, chunks, unparsed, rule_cards, station_index, amenities)
    write_docs(catalog, manifest, unparsed)
    print(json.dumps({
        "catalog": catalog["summary"],
        "sourceStatus": catalog["sourceStatus"],
        "sourceUsefulness": catalog["sourceUsefulness"],
        "docs": "docs/knowledge_database.md"
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
